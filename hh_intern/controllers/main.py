# -*- coding: utf-8 -*-
from odoo import http
from odoo.http import request
from odoo.addons.web.controllers.main import serialize_exception,content_disposition
from docx.enum.shape import WD_INLINE_SHAPE

from io import BytesIO, StringIO
from docxtpl import DocxTemplate, InlineImage, CheckedBox
from docx.shared import Mm, Inches, Pt
# from docx import Document
from tempfile import NamedTemporaryFile
import os

from odoo.addons.hh_intern.models import intern_utils
import zipfile

import logging
_logger = logging.getLogger(__name__)

class CreateDocument(http.Controller):

    @http.route('/web/binary/download_document_hoso', type='http', auth="public")
    def download_document_hoso(self, model, id, filename=None, **kwargs):
        invoice = request.env[model].browse(int(id))
        document = request.env['intern.document'].search([('name', '=', 'CV')], limit=1)
        reponds = BytesIO()
        archive = zipfile.ZipFile(reponds, 'w', zipfile.ZIP_DEFLATED)
        finalDoc = invoice.createHeaderDoc()
        if finalDoc is not None:
            archive.write(finalDoc.name,u"名簿リスト.docx")
            os.unlink(finalDoc.name)
        else:
            return

        listtmp = sorted(invoice.interns_exam_doc, key=lambda x: x.sequence_exam)

        for i, intern in enumerate(listtmp):
            if intern.pass_exam:
                childDoc = invoice.createCVDoc(document[0], intern, i)
                archive.write(childDoc.name, 'cv_%d_%s.docx' % ((i + 1), intern_utils.name_with_underscore(intern.name)))
                os.unlink(childDoc.name)

        archive.close()
        reponds.flush()
        ret_zip = reponds.getvalue()
        reponds.close()

        return request.make_response(ret_zip,
                                     [('Content-Type', 'application/zip'),
                                      ('Content-Disposition', content_disposition(filename))])


    @http.route('/web/binary/download_document', type='http', auth="public")
    def download_cv(self, model, id, filename=None, **kwargs):
        # invoice = request.env[model].search([('id', '=', id)])
        invoice = request.env[model].browse(int(id))
        # request._cr.execute('SELECT intern_id FROM intern_order WHERE intern_order.invoice_id = %s' % id)
        # tmpresult = request._cr.dictfetchall()
        # if len(tmpresult) == 1:
        #     ids = tmpresult[0]['intern_id']
        #     if len(ids) == len(invoice.interns):
        #         interns = request.env['intern.intern'].browse(ids)
        #         invoice.interns = interns

        document = request.env['intern.document'].search([('name', '=', 'CV')], limit=1)
        finalDoc = invoice.createHeaderDoc()
        # byteIo = BytesIO()
        # finalDoc.save(byteIo)
        # byteIo.seek(0)
        reponds = BytesIO()
        archive = zipfile.ZipFile(reponds, 'w', zipfile.ZIP_DEFLATED)
        # merge_doc = None
        if finalDoc is not None:
            archive.write(finalDoc.name,u"名簿リスト.docx")

            # merge_doc = Document(finalDoc.name)
            os.unlink(finalDoc.name)
        else:
            return

        listtmp = None
        # if invoice.order:
        #     ids = []
        #     for intern in invoice.interns:
        #         ids.append(intern.id)
        #
        #     try:
        #         listtmp = request.env['intern.intern'].search([('id', 'in', ids)], order="%s" % (invoice.order))
        #     except:
        #         listtmp = invoice.interns
        # else:

        listtmp = sorted(invoice.interns_exam_doc,key=lambda x: x.sequence_exam)

        for i, intern in enumerate(listtmp):
            childDoc = invoice.createCVDoc(document[0], intern, i)
            archive.write(childDoc.name,'cv_%d_%s.docx'%((i+1),intern_utils.name_with_underscore(intern.name)))

            # tmpDoc = Document(childDoc.name)
            # for element in tmpDoc.element.body:
            #     merge_doc.element.body.append(element)

            os.unlink(childDoc.name)

        # tempFile = NamedTemporaryFile(delete=False)
        # merge_doc.save(tempFile.name)
        # archive.write(tempFile.name,"full.docx")
        # os.unlink(tempFile.name)

        archive.close()
        reponds.flush()
        ret_zip = reponds.getvalue()
        reponds.close()

        return request.make_response(ret_zip,
                                         [('Content-Type', 'application/zip'),
                                          ('Content-Disposition', content_disposition(filename))])

    @http.route('/web/binary/download_proletter_document', type='http', auth="public")
    def download_proletter_document(self, model, id, enterprise, filename, **kwargs):
        # invoices = request.env[model].search([('id', '=', int(id))])
        # invoice = invoices[0]
        invoice = request.env[model].browse(int(id))
        # request._cr.execute('SELECT intern_id FROM internpass_order WHERE internpass_order.invoice_id = %s' % id)
        # tmpresult = request._cr.dictfetchall()
        # if len(tmpresult) == 1:
        #     ids = tmpresult[0]['intern_id']
        #     if len(ids) == len(invoice.interns_pass):
        #         interns_pass = request.env['intern.intern'].browse(ids)
        #         invoice.interns_pass = interns_pass


        reponds = BytesIO()
        archive = zipfile.ZipFile(reponds, 'w', zipfile.ZIP_DEFLATED)

        doc_list_send = invoice.create_list_of_sent_en(enterprise)
        archive.write(doc_list_send.name, u'推薦書 - ENG.docx')
        os.unlink(doc_list_send.name)

        doc_list_send_jp = invoice.create_list_of_sent_jp(enterprise)
        archive.write(doc_list_send_jp.name, u'推薦書.docx')
        os.unlink(doc_list_send_jp.name)
        archive.close()
        reponds.flush()
        ret_zip = reponds.getvalue()
        reponds.close()

        return request.make_response(ret_zip,
                                     [('Content-Type', 'application/zip'),
                                      ('Content-Disposition', content_disposition(filename))])

    @http.route('/web/binary/download_extern_document_specific', type='http', auth="public")
    def download_extern_document_specific(self, model, id, document,enterprise,filename=None, **kwargs):
        # invoice = request.env[model].search([('id', '=', id)])
        invoice = request.env[model].browse(int(id))
        # request._cr.execute('SELECT intern_id FROM internpass_order WHERE internpass_order.invoice_id = %s' % id)
        # tmpresult = request._cr.dictfetchall()
        # if len(tmpresult) == 1:
        #     ids = tmpresult[0]['intern_id']
        #     if len(ids) == len(invoice.interns_pass):
        #         interns_pass = request.env['intern.intern'].browse(ids)
        #         invoice.interns_pass = interns_pass


        # if invoice.dispatchcom2 and len(invoice.dispatchcom2) >1:
            # request._cr.execute('SELECT dispatch_id FROM dispatch_order WHERE dispatch_order.invoice_id = %s' % id)
            # tmpresult2 = request._cr.dictfetchall()
            # if len(tmpresult2) ==1:
            #     ids2 = tmpresult2[0]['dispatch_id']
            #     if len(ids2) ==  len(invoice.dispatchcom2):
            #         dispatchs = request.env['dispatchcom2'].browse(ids2)
            #         invoice.dispatchcom2 = dispatchs

        reponds = BytesIO()
        archive = zipfile.ZipFile(reponds, 'w', zipfile.ZIP_DEFLATED)

        list_interns = invoice.interns_pass_doc
        if invoice.hoso_created:
            list_interns = invoice.interns_pass_doc_hs

        interns_pass = sorted(list_interns,key=lambda x: x.sequence_pass)
        enterprise_id = int(enterprise)

        if document == 'Doc1-3':
            counter = 0
            for i, intern in enumerate(interns_pass):
                if intern.enterprise.id != enterprise_id:
                    continue
                doc1_3 = invoice.create_doc_1_3(intern, counter)
                archive.write(doc1_3.name, '1_3_%d_%s.docx' % ((counter + 1), intern_utils.name_with_underscore(intern.name)))
                os.unlink(doc1_3.name)
                counter += 1

        elif document == 'Doc1-10':
            counter = 0
            for i, intern in enumerate(interns_pass):
                if intern.enterprise.id != enterprise_id:
                    continue
                counter += 1
                doc1_10 = invoice.create_doc_1_10(intern)
                archive.write(doc1_10.name,
                              '1_10_%d_%s.docx' % (counter, intern_utils.name_with_underscore(intern.name)))
                os.unlink(doc1_10.name)
        elif document == 'Doc1-13':
            doc1_13_1 = invoice.create_1_13_1()

            archive.write(doc1_13_1.name, u'1-13号 HOANG HUNG JAPAN 訓連センター.docx')
            os.unlink(doc1_13_1.name)

            doc1_13_2 = invoice.create_1_13_2()
            archive.write(doc1_13_2.name, u'1-13号HOANG HUNG 会社.docx')
            os.unlink(doc1_13_2.name)
        elif document == 'Doc1-20':
            doc1_20 = invoice.create_doc_1_20()
            archive.write(doc1_20.name, '1_20.docx')
            os.unlink(doc1_20.name)
        elif document == 'Doc1-21':
            counter=0
            for i, intern in enumerate(interns_pass):
                if intern.enterprise.id != enterprise_id:
                    continue
                counter+=1
                doc1_21 = invoice.create_doc_1_21(intern)
                archive.write(doc1_21.name, '1_21_%d_%s.docx' % (counter, intern_utils.name_with_underscore(intern.name)))
                os.unlink(doc1_21.name)
        elif document == 'Doc1-28':
            counter = 0
            for i, intern in enumerate(interns_pass):
                if intern.enterprise.id != enterprise_id:
                    continue
                doc1_28 = invoice.create_doc_1_28(intern,counter)
                archive.write(doc1_28.name, '1_28_%d_%s.docx' % ((counter + 1), intern_utils.name_with_underscore(intern.name)))
                os.unlink(doc1_28.name)
                counter +=1
        elif document == 'Doc1-29':
            doc1_29 = invoice.create_doc_1_29(enterprise_id)
            archive.write(doc1_29.name, '1_29.docx')
            os.unlink(doc1_29.name)
        elif document == 'DocCCDT':
            docCCDT = invoice.create_certification_end_train(enterprise_id)
            archive.write(docCCDT.name, 'CCDT.docx')
            os.unlink(docCCDT.name)
        elif document == 'HDPC':
            counter=0
            for i, intern in enumerate(interns_pass):
                if intern.enterprise.id != enterprise_id:
                    continue
                counter+=1
                hdtn = invoice.create_hdtn(intern)
                archive.write(hdtn.name, 'hdtn_%d_%s.docx' % (counter, intern_utils.name_with_underscore(intern.name)))
                os.unlink(hdtn.name)

                hdtv = invoice.create_hdtv(intern)
                archive.write(hdtv.name, 'hdtv_%d_%s.docx' % (counter, intern_utils.name_with_underscore(intern.name)))
                os.unlink(hdtv.name)
        elif document == 'PROLETTER':
            doc_list_send = invoice.create_list_of_sent_en(enterprise_id)
            archive.write(doc_list_send.name, u'推薦書 - ENG.docx')
            os.unlink(doc_list_send.name)

            doc_list_send_jp = invoice.create_list_of_sent_jp(enterprise_id)
            archive.write(doc_list_send_jp.name, u'推薦書.docx')
            os.unlink(doc_list_send_jp.name)

        elif document == 'VISAFORM':
            counter = 0
            for i, intern in enumerate(interns_pass):
                if intern.enterprise.id != enterprise_id:
                    continue
                counter+=1
                visa_form = invoice.create_visa_application_form(intern)
                archive.write(visa_form.name, 'hdtn_%d_%s.docx' % (counter, intern_utils.name_with_underscore(intern.name)))
                os.unlink(visa_form.name)

                hdtv = invoice.create_hdtv(intern)
                archive.write(hdtv.name, 'hdtv_%d_%s.docx' % (counter, intern_utils.name_with_underscore(intern.name)))
                os.unlink(hdtv.name)
        elif document == 'DSLD':
            reponds.close()
            return request.make_response(invoice.create_danh_sach_lao_dong(enterprise_id),
                                                  headers=[('Content-Disposition',
                                                            content_disposition('DSLD.docx')),
                                                           ('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')])

        elif document == 'CheckList':
            reponds.close()
            return request.make_response(invoice.create_check_list(enterprise_id),headers=[('Content-Disposition',
                                                            content_disposition('Check_List.docx')),
                                                           ('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')])

        elif document == 'Doc4-8':
            reponds.close()
            return request.make_response(invoice.create_48(enterprise_id), headers=[('Content-Disposition',
                                                                                             content_disposition(
                                                                                                 '4-8-入国前講習実施記録 (5-10).docx')),
                                                                                            ('Content-Type',
                                                                                             'application/vnd.openxmlformats-officedocument.wordprocessingml.document')])
        elif document == 'Master':
            master = invoice.create_master(enterprise_id)
            archive.write(master.name, 'Master.docx')
            os.unlink(master.name)
        elif document == 'Doc1-27':
            reponds.close()
            return request.make_response(invoice.create_1_27(enterprise_id), headers=[('Content-Disposition',
                                                                                             content_disposition(
                                                                                                 'Doc1-27.docx')),
                                                                                            ('Content-Type',
                                                                                             'application/vnd.openxmlformats-officedocument.wordprocessingml.document')])

        archive.close()
        reponds.flush()
        ret_zip = reponds.getvalue()
        reponds.close()

        return request.make_response(ret_zip,
                                     [('Content-Type', 'application/zip'),
                                      ('Content-Disposition', content_disposition(filename))])

    @http.route('/web/binary/download_extern_document', type='http', auth="public")
    def download_extern_document(self, model, id, enterprise, filename=None, **kwargs):
        # invoice = request.env[model].search([('id', '=', id)])
        invoice = request.env[model].browse(int(id))
        enterprise_id = int(enterprise)

        list_interns = invoice.interns_pass_doc
        if invoice.hoso_created:
            list_interns = invoice.interns_pass_doc_hs

        interns_pass = sorted(list_interns, key=lambda x: x.sequence_pass)
        # request._cr.execute('SELECT intern_id FROM internpass_order WHERE internpass_order.invoice_id = %s' % id)
        # tmpresult = request._cr.dictfetchall()
        # if len(tmpresult) == 1:
        #     ids = tmpresult[0]['intern_id']
        #     if len(ids) == len(invoice.interns_pass):
        #         interns_pass = request.env['intern.intern'].browse(ids)
        #         invoice.interns_pass = interns_pass
        #
        # if invoice.dispatchcom2 and len(invoice.dispatchcom2) >1:
        #     request._cr.execute('SELECT dispatch_id FROM dispatch_order WHERE dispatch_order.invoice_id = %s' % id)
        #     tmpresult2 = request._cr.dictfetchall()
        #     if len(tmpresult2)==1:
        #         ids2 = tmpresult2[0]['dispatch_id']
        #         if len(ids2) == len(invoice.dispatchcom2):
        #             dispatchs = request.env['dispatchcom2'].browse(ids2)
        #             invoice.dispatchcom2 = dispatchs


        reponds = BytesIO()
        archive = zipfile.ZipFile(reponds, 'w', zipfile.ZIP_DEFLATED)

        # checklist = request.env['intern.document'].search([('name', '=', "Checklist")], limit=1)
        # if checklist:
        #     stream = BytesIO(checklist[0].attachment.decode("base64"))
        #     tpl = DocxTemplate(stream)
        #     tempFile = NamedTemporaryFile(delete=False)
        #     tpl.render({})
        #     tpl.save(tempFile)
        #     tempFile.flush()
        #     tempFile.close()
        #     archive.write(tempFile.name, 'Checklist.docx')
        #     os.unlink(tempFile.name)

        file_maps = {}

        doc1_13_1 = invoice.create_1_13_1()

        file_maps.update({u'1-13号 HOANG HUNG JAPAN 訓連センター.docx':doc1_13_1.name})
        # archive.write(doc1_13_1.name, u'1-13号 HOANG HUNG JAPAN 訓連センター.docx')
        # os.unlink(doc1_13_1.name)

        doc1_13_2 = invoice.create_1_13_2()
        file_maps.update({u'1-13号HOANG HUNG 会社.docx': doc1_13_2.name})
        # archive.write(doc1_13_2.name, u'1-13号HOANG HUNG 会社.docx')
        # os.unlink(doc1_13_2.name)

        # master = invoice.create_master(enterprise_id)
        # file_maps.update({'Master.docx': master.name})
        # archive.write(master.name, 'Master.docx')
        # os.unlink(master.name)

        doc1_29 = invoice.create_doc_1_29(enterprise_id)
        file_maps.update({'1_29.docx': doc1_29.name})
        # archive.write(doc1_29.name, '1_29.docx')
        # os.unlink(doc1_29.name)

        doc_list_send = invoice.create_list_of_sent_en(enterprise_id)
        file_maps.update({u'推薦書 - ENG.docx': doc_list_send.name})
        # archive.write(doc_list_send.name, u'推薦書 - ENG.docx')
        # os.unlink(doc_list_send.name)

        doc_list_send_jp = invoice.create_list_of_sent_jp(enterprise_id)
        file_maps.update({u'推薦書.docx': doc_list_send_jp.name})
        # archive.write(doc_list_send_jp.name, u'推薦書.docx')
        # os.unlink(doc_list_send_jp.name)

        doc1_20 = invoice.create_doc_1_20()
        file_maps.update({'1_20.docx': doc1_20.name})
        # archive.write(doc1_20.name, '1_20.docx')
        # os.unlink(doc1_20.name)

        docCCDT = invoice.create_certification_end_train(enterprise_id)
        file_maps.update({u'事前講習実施報告書.docx': docCCDT.name})
        # archive.write(docCCDT.name, u'事前講習実施報告書.docx')
        # os.unlink(docCCDT.name)

        iterate_intern = 0
        for i, intern in enumerate(interns_pass):
            if intern.enterprise.id != enterprise_id:
                continue

            doc1_3 = invoice.create_doc_1_3(intern, iterate_intern)
            file_maps.update({'1_3_%d_%s.docx' % ((iterate_intern+1),intern_utils.name_with_underscore(intern.name)): doc1_3.name})
            # archive.write(doc1_3.name, '1_3_%d_%s.docx' % ((iterate_intern+1),intern_utils.name_with_underscore(intern.name)))
            # os.unlink(doc1_3.name)

            doc1_10 = invoice.create_doc_1_10(intern)
            file_maps.update(
                {'1_10_%d_%s.docx' % ((iterate_intern+1),intern_utils.name_with_underscore(intern.name)): doc1_10.name})
            # archive.write(doc1_10.name, '1_10_%d_%s.docx' % ((iterate_intern+1),intern_utils.name_with_underscore(intern.name)))
            # os.unlink(doc1_10.name)



            doc1_21 = invoice.create_doc_1_21(intern)
            file_maps.update(
                {'1_21_%d_%s.docx' % ((iterate_intern+1),intern_utils.name_with_underscore(intern.name)): doc1_21.name})
            # archive.write(doc1_21.name, '1_21_%d_%s.docx' % ((iterate_intern+1),intern_utils.name_with_underscore(intern.name)))
            # os.unlink(doc1_21.name)

            doc1_28 = invoice.create_doc_1_28(intern,i)
            file_maps.update(
                {'1_28_%d_%s.docx' % ((iterate_intern+1),intern_utils.name_with_underscore(intern.name)): doc1_28.name})
            # archive.write(doc1_28.name, '1_28_%d_%s.docx' % ((iterate_intern+1),intern_utils.name_with_underscore(intern.name)))
            # os.unlink(doc1_28.name)

            hdtn = invoice.create_hdtn(intern)
            file_maps.update(
                {'hdtn_%d_%s.docx' % ((iterate_intern+1),intern_utils.name_with_underscore(intern.name)): hdtn.name})
            # archive.write(hdtn.name, 'hdtn_%d_%s.docx' % ((iterate_intern+1),intern_utils.name_with_underscore(intern.name)))
            # os.unlink(hdtn.name)

            hdtv = invoice.create_hdtv(intern)
            file_maps.update(
                {'hdtv_%d_%s.docx' % ((iterate_intern+1),intern_utils.name_with_underscore(intern.name)): hdtv.name})
            # archive.write(hdtv.name, 'hdtv_%d_%s.docx' % ((iterate_intern+1),intern_utils.name_with_underscore(intern.name)))
            # os.unlink(hdtv.name)

            iterate_intern+=1

        for key in file_maps:
            archive.write(file_maps[key],key)
            os.unlink(file_maps[key])

        archive.close()
        reponds.flush()
        ret_zip = reponds.getvalue()
        reponds.close()

        return request.make_response(ret_zip,
                                     [('Content-Type', 'application/zip'),
                                      ('Content-Disposition', content_disposition(filename))])


