# -*- coding: utf-8 -*-
from odoo import models, fields, api

from tempfile import TemporaryFile, NamedTemporaryFile
from io import BytesIO
from datetime import datetime, date
from docxtpl import DocxTemplate, InlineImage, CheckedBox, CheckBox, RichText, Tick
from docx.shared import Mm, Inches
from docx import Document
import intern_utils
from odoo import SUPERUSER_ID
import os
from odoo.exceptions import UserError, ValidationError
from StringIO import StringIO
import logging
from dateutil.relativedelta import relativedelta

_logger = logging.getLogger(__name__)

class Invoice(models.Model):
    _name = 'intern.invoice'
    _description = u'Đơn hàng'

    _order = 'id desc'

    name = fields.Char("Tên đơn hàng", required=True)
    name_of_guild = fields.Char("Tên nghiệp đoàn")
    enterprise = fields.Char("Tên xí nghiệp")
    order = fields.Char()

    place_to_work = fields.Char("Địa điểm làm việc")
    legal_name = fields.Selection([('chauhung', 'Châu Hưng'), ('hoanghung', 'Hoàng Hưng'), ('tracodi', 'Tracodi'),('thuanan','Thuận An')],
                                  'Pháp nhân')

    room_pttt = fields.Many2one('department',string='Phòng PTTT')
    employee_pttt = fields.Many2one('hh.employee', string=u"Cán bộ PTTT")


    document = fields.Selection([('Doc1-3', '1-3'), ('Doc1-10', '1-10'), ('Doc1-13', '1-13'),('Doc1-20', '1-20'),
                                 ('Doc1-21', '1-21'),('Doc1-28', '1-28'),('Doc1-29', '1-29'),('DocCCDT','Chứng chỉ kết thúc Đào tạo'),
                                 ('HDPC', 'Hợp đồng PC'),('PROLETTER', 'Thư tiến cử'),('DSLD','Danh sách lao động'),('Doc4-8','4-8')], string='Hồ sơ in',store=True
                                    )


    interns_clone = fields.One2many('intern.internclone','invoice_id',string='Thực tập sinh')

    interns_promoted = fields.One2many('intern.internclone', 'invoice_id')
    interns_confirm_exam = fields.One2many('intern.internclone', 'invoice_id')
    interns_escape_exam = fields.One2many('intern.internclone', 'invoice_id')
    interns_pass_new = fields.One2many('intern.internclone', 'invoice_id')
    interns_preparatory = fields.One2many('intern.internclone', 'invoice_id')
    interns_cancel_pass = fields.One2many('intern.internclone', 'invoice_id')
    interns_departure = fields.One2many('intern.internclone', 'invoice_id')


    @api.multi
    def create_proletter_doc(self,enterprise):

        list_interns = self.interns_pass_doc
        if self.hoso_created:
            list_interns = self.interns_pass_doc_hs

        interns_pass = sorted(list_interns, key=lambda x: x.sequence_pass)

        if interns_pass is None or len(interns_pass) == 0:
            raise ValidationError("Không có thực tập sinh nào trong danh sách trúng tuyển")

        error2 = ""
        if not self.dispatchcom1:
            error2 = error2 + u"- pháp nhân\n"
        if not self.guild:
            error2 = error2 + u"- nghiệp đoàn\n"
        # if not self.enterprise_doc:
        #     error2 = error2 + u"- xí nghiệp\n"
        if self.year_expire==0:
            error2 = error2 + u"- thời hạn hợp đồng\n"

        if not self.person_sign_proletter\
                or not self.position_person_sign:
            error2 = error2+u"- người ký thư tiến cử\n"
        if not self.day_create_letter_promotion or not self.month_create_letter_promotion or not self.year_create_letter_promotion:
            error2 = error2 + u"- ngày làm thư tiến cử\n"
        if not self.job_en or not self.job_jp or not self.job_vi:
            error2 = error2 + u"- ngành nghề xin thư tiến cử\n"

        if not self.month_departure_doc or not self.year_departure_doc:
            error2 = error2 + u"- ngày nộp xuất cảnh dự kiến\n"

        if error2:
            raise ValidationError(u"Thiếu thông tin bổ sung cho hồ sơ: \n%s" % error2)


        if not self.dispatchcom1.name_jp or not self.dispatchcom1.name_en or not self.dispatchcom1.name \
            or not self.dispatchcom1.director or not self.dispatchcom1.position_director or not self.dispatchcom1.position_director_vi \
            or not self.dispatchcom1.address_vi \
                or not self.dispatchcom1.address_en \
            or not self.dispatchcom1.phone_number:

            raise ValidationError(u"Thiếu thông tin của pháp nhân")

        enterprise_obj = self.env['intern.enterprise'].browse(enterprise)
        return {
            'type': 'ir.actions.act_url',
            'url': '/web/binary/download_proletter_document?model=intern.invoice&id=%s&enterprise=%s&filename=%s_%s_ThuTienCu.zip' % (
                self.id,enterprise,enterprise_obj.name_romaji, self.guild.name_acronym),
            'target': 'self', }


    @api.multi
    def create_extern_doc(self, enterprise_id, document):

        if document and document == 'PROLETTER':
            return self.create_proletter_doc(enterprise_id)

        # if self.interns_pass_doc is None and  or len(self.interns_pass_doc) is 0:
        list_interns = self.interns_pass_doc
        if self.hoso_created:
            list_interns = self.interns_pass_doc_hs
        if list_interns is None or len(list_interns) == 0:
            raise ValidationError("Không có thực tập sinh nào trong danh sách trúng tuyển")

        # Validate for doc 1-3
        count_intern_by_enterprise = 0
        for itern in list_interns:
            if itern.enterprise.id != enterprise_id:
                continue
            count_intern_by_enterprise += 1

        if document and document == 'Doc1-3':
            for itern in list_interns:
                if itern.enterprise.id != enterprise_id:
                    continue
                error = ""
                if not itern.name:
                    error = error + u'- tên tiếng Việt\n'
                if not itern.name_in_japan:
                    error = error + u'- tên tiếng Nhật\n'
                if not itern.day or not itern.month or not itern.year:
                    error = error + u'- ngày sinh\n'
                if not itern.gender:
                    error = error + u'- giới tính\n'
                if not itern.marital_status:
                    error = error + u'- tình trạng hôn nhân\n'
                if not itern.hktt:
                    error = error + u'- Địa chỉ HKTT\n'
                if not itern.last_education_from_month or not itern.last_education_from_year \
                        or not itern.last_education_to_month or not itern.last_education_to_year or not itern.last_school_education_jp:
                    error = error + u'- Trường học mới nhất\n'

                if not itern.time_start_at_pc_from_month or not itern.time_start_at_pc_from_year \
                        or (itern.time_at_pc_year == 0 and itern.time_at_pc_month == 0):
                    error = error + u'- thông tin làm việc tại Công ty PC2\n'
                if not itern.dispatchcom2:
                    error = error + u"- công ty phái cử 2\n"
                if error:
                    raise ValidationError(u"Thiếu thông tin của %s:\n%s" % (itern.name, error))

            error2 = ""
            if not self.day_create_letter_promotion or not self.month_create_letter_promotion or not self.year_create_letter_promotion:
                error2 = error2 + u"- ngày làm thư tiến cử\n"

            if error2:
                raise ValidationError(u"Thiếu thông tin bổ sung cho hồ sơ: \n%s" % error2)
        elif document and document == 'DocCCDT': #validate CCDT

            error2 = ""
            if not self.day_create_plan_training or not self.month_create_plan_training or not self.year_create_plan_training:
                error2 = error2 + u"- ngày lập kế hoạch đào tạo\n"
            if not self.day_start_training or not self.month_start_training or not self.year_start_training:
                error2 = error2 + u"- ngày bắt đầu khoá học\n"
            if not self.day_end_training or not self.month_end_training or not self.year_end_training:
                error2 = error2 + u"- ngày kết thúc khoá học\n"
            if not self.date_create_plan_training_report_customer or not self.month_create_plan_training_report_customer or not self.year_create_plan_training_report_customer:
                error2 = error2 + u"- ngày lập kế hoạch đào tạo báo cáo khách hàng\n"

            if not self.length_training or not self.hours_training:
                error2 = error2 + u"- thời gian đào tạo khoá học\n"
            if not self.training_center:
                error2 = error2 + u"- trung tâm đào tạo\n"
            if not self.guild:
                error2 = error2 + u"- nghiệp đoàn\n"

            if error2:
                raise ValidationError(u"Thiếu thông tin bổ sung cho hồ sơ: \n%s"%error2)
        elif document and document == 'DSLD':
            error2 = ""
            if not self.job_en:
                error2 = error2 + u"- ngành nghề xin thư tiến cử TA\n"
            if not self.month_departure_doc or not self.year_departure_doc:
                error2 = error2 + u"- thời gian xuất cảnh\n"

            if error2:
                raise ValidationError(u"Thiếu thông tin: \n%s"%error2)
            error = ""
            for itern in list_interns:
                if not itern.hktt:
                    error = error + u'- Địa chỉ HKTT\n'
                if error:
                    raise ValidationError(u"Thiếu thông tin của %s:\n%s" % (itern.name,error))
        elif document and document == 'CheckList':
            _logger.info('AAAAA')
        elif document and document == 'Doc4-8':
            error2 = ""
            if not self.day_start_training or not self.month_start_training or not self.year_start_training:
                error2 = error2 + u"- ngày bắt đầu khoá học\n"
            if not self.day_end_training or not self.month_end_training or not self.year_end_training:
                error2 = error2 + u"- ngày kết thúc khoá học\n"
            if not self.training_center:
                error2 = error2 + u"- trung tâm đào tạo\n"
            if not self.date_departure_doc:
                error2 = error2 + u"- ngày xuất cảnh dự kiến\n"
            if error2:
                raise ValidationError(u"Thiếu thông tin: \n%s" % error2)
        elif document and document == 'Doc1-27':
            error2 = ""
            if not self.day_create_letter_promotion or not self.month_create_letter_promotion or not self.year_create_letter_promotion:
                error2 = error2 + u"- ngày làm thư tiến cử\n"

            if error2:
                raise ValidationError(u"Thiếu thông tin bổ sung cho hồ sơ: \n%s" % error2)
        else:
            error2 = ""
            if not self.person_sign_proletter \
                    or not self.position_person_sign:
                error2 = error2 + u"- người ký thư tiến cử\n"
            if not self.day_create_letter_promotion or not self.month_create_letter_promotion or not self.year_create_letter_promotion:
                error2 = error2 + u"- ngày làm thư tiến cử\n"

            if not self.day_create_plan_training or not self.month_create_plan_training or not self.year_create_plan_training:
                error2 = error2 + u"- ngày lập kế hoạch đào tạo\n"
            if not self.day_start_training or not self.month_start_training or not self.year_start_training:
                error2 = error2 + u"- ngày bắt đầu khoá học\n"
            if not self.day_end_training or not self.month_end_training or not self.year_end_training:
                error2 = error2 + u"- ngày kết thúc khoá học\n"
            if not self.date_create_plan_training_report_customer or not self.month_create_plan_training_report_customer or not self.year_create_plan_training_report_customer:
                error2 = error2 + u"- ngày lập kế hoạch đào tạo báo cáo khách hàng\n"
            if not self.day_pay_finance1 or not self.month_pay_finance1 or not self.year_pay_finance1:
                error2 = error2 + u"- ngày nộp tài chính lần 1\n"
            if not self.day_pay_finance2 or not self.month_pay_finance2 or not self.year_pay_finance2:
                error2 = error2 + u"- ngày nộp tài chính lần 2\n"
            if not self.month_departure_doc or not self.year_departure_doc:
                error2 = error2 + u"- ngày nộp xuất cảnh dự kiến\n"

            if not self.length_training or not self.hours_training:
                error2 = error2 + u"- thời gian đào tạo khoá học\n"
            if not self.training_center:
                error2 = error2 + u"- trung tâm đào tạo\n"
            if not self.guild:
                error2 = error2 + u"- nghiệp đoàn\n"

            if not self.dispatchcom1:
                error2 = error2 + u"- pháp nhân\n"
            if not self.dispatchcom2:
                error2 = error2 + u"- công ty phái cử 2\n"

            if not self.name_working_department:
                error2 = error2 + u"- bộ phận TTS sẽ làm việc\n"
            if not self.job_en or not self.job_jp or not self.job_vi:
                error2 = error2 + u"- ngành nghề xin thư tiến cử\n"

            if error2:
                raise ValidationError(u"Thiếu thông tin bổ sung cho hồ sơ: \n%s" % error2)

            for itern in list_interns:
                if itern.enterprise.id != enterprise_id:
                    continue
                error = ""
                if not itern.name:
                    error = error+u'- tên tiếng Việt\n'
                if not itern.name_in_japan:
                    error = error + u'- tên tiếng Nhật\n'
                if not itern.day or not itern.month or not itern.year:
                    error = error + u'- ngày sinh\n'
                if not itern.gender:
                    error = error + u'- giới tính\n'
                if not itern.marital_status:
                    error = error + u'- tình trạng hôn nhân\n'
                if (not itern.identity and not itern.identity_2):
                    error = error + u'- CMND\n'
                if itern.identity and not itern.place_cmnd:
                    error = error + u'- nơi cấp CMND\n'
                if not itern.day_identity or not itern.month_identity or not itern.year_identity:
                    error = error + u'- ngày cấp CMND\n'
                if not itern.hktt:
                    error = error + u'- Địa chỉ HKTT\n'
                if not itern.last_education_from_month or not itern.last_education_from_year \
                    or not itern.last_education_to_month or not itern.last_education_to_year or not itern.last_school_education_jp:
                    error = error + u'- Trường học mới nhất\n'

                if not itern.time_start_at_pc_from_month or not itern.time_start_at_pc_from_year  \
                    or (itern.time_at_pc_year == 0 and itern.time_at_pc_month == 0):
                    error = error + u'- thông tin làm việc tại Công ty PC2\n'

                if not itern.contact_person or not itern.contact_address or \
                        not itern.contact_relative or not itern.contact_phone:
                    error = error + u'- người thân khi cần liên lạc\n'

                if not itern.enterprise:
                    error = error + u"- xí nghiệp\n"

                if error:
                    raise ValidationError(u"Thiếu thông tin của %s:\n%s" % (itern.name,error))



            if not self.dispatchcom1.name_jp or not self.dispatchcom1.name_en or not self.dispatchcom1.name \
                or not self.dispatchcom1.director or not self.dispatchcom1.position_director or not self.dispatchcom1.position_director_vi \
                or not self.dispatchcom1.address_vi \
                    or not self.dispatchcom1.address_en \
                or not self.dispatchcom1.phone_number or not self.dispatchcom1.date_create or not self.dispatchcom1.mission \
                or not self.dispatchcom1.capital or not self.dispatchcom1.revenue:
                raise ValidationError(u"Thiếu thông tin của pháp nhân")


            if not self.training_center.name_jp or not self.training_center.address_en \
                    or not self.training_center.date_create or not self.training_center.phone_number \
                or not self.training_center.responsive_person or not self.training_center.mission:
                raise ValidationError(u"Thiếu thông tin của trung tâm đào tạo")
            if not self.guild.name_in_en or not self.guild.address_in_romaji or \
                not self.guild.position_of_responsive_vi or not self.guild.name_of_responsive_romaji:
                raise ValidationError(u"Thiếu thông tin của nghiệp đoàn")



        enterprise_obj = self.env['intern.enterprise'].browse(enterprise_id)

        if document and document !='All':
            return {
                'type': 'ir.actions.act_url',
                'url': '/web/binary/download_extern_document_specific?model=intern.invoice&id=%s&document=%s&enterprise=%s&filename=%s_%s_%s.zip' % (
                    self.id, document,enterprise_id, enterprise_obj.name_romaji, self.guild.name_acronym, document),
                'target': 'self', }
        else:
            return {
                'type': 'ir.actions.act_url',
                'url': '/web/binary/download_extern_document?model=intern.invoice&id=%s&enterprise=%s&filename=%s_%s_HoSo.zip' % (
                    self.id,enterprise_id,enterprise_obj.name_romaji,self.guild.name_acronym),
                'target': 'self', }


    @api.multi
    def create_doc(self):
        _logger.info("CREATE DOC")

        if self.interns_exam_doc is None or len(self.interns) is 0:
            raise ValidationError("Không có thực tập sinh nào trong danh sách thi tuyển")

        for intern in self.interns_exam_doc:
            if intern.date_of_birth is None:
                raise ValidationError(u"Thiếu thông tin %s của %s"%(u'ngày sinh',intern.name))

            if not intern.name_in_japan:
                raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Tên tiếng Nhật', intern.name))
            if not intern.gender:
                raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Giới tính', intern.name))
            if not intern.blood_group:
                raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Nhóm máu', intern.name))
            if intern.vision_left is None or intern.vision_right is None:
                raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Thị lực', intern.name))
            if not intern.iq_percentage or intern.iq_percentage is None or intern.iq_percentage is 0:
                raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Điểm IQ', intern.name))
            if intern.check_kureperin is None:
                raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Điểm cộng dồn', intern.name))
            if not intern.marital_status:
                raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Tình trạng hôn nhân', intern.name))
            if not intern.province or not intern.address:
                raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Quê quán', intern.name))
            if not intern.certification:
                raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Bằng cấp', intern.name))
            if intern.certification.id >2 and not intern.specialized:
                raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Chuyên ngành', intern.name))

            if intern.height is None:
                raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Chiều cao', intern.name))
            if intern.weight is None:
                raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Cân nặng', intern.name))
            if not intern.favourite or not intern.strong or \
                    not intern.weak or not intern.family_income or not intern.motivation or \
                    not intern.income_after_three_year or not intern.job_after_return or \
                    not intern.prefer_object or not intern.memory or not intern.valuable:
                raise ValidationError(u"Thiếu thông tin %s của %s" % (u'bổ sung tiếng Nhật', intern.name))

            if intern.educations is None or len(intern.educations) is 0:
                raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Lý lịch Học tập', intern.name))

            if intern.family_members is None or len(intern.family_members) is 0:
                raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Thành viên gia đình', intern.name))

        return {
            'type': 'ir.actions.act_url',
            'url': '/web/binary/download_document?model=intern.invoice&id=%s&filename=%s.zip' %(str(self.id) ,self.name),
            'target': 'self',}


    @api.multi
    def create_doc_hoso(self):
        return {
            'type': 'ir.actions.act_url',
            'url': '/web/binary/download_document_hoso?model=intern.invoice&id=%s&filename=%s.zip' % (
            str(self.id), self.name),
            'target': 'self', }

    def createHeaderDoc(self):
        docs = self.env['intern.document'].search([('name', '=', "CV_HEAD")], limit=1)
        logo = self.env['intern.document'].search([('name', '=', "Logo")], limit=1)
        streamDoc = BytesIO(docs[0].attachment.decode("base64"))
        target_document = None
        tempFile1 = None
        if len(self.interns_exam_doc)>20:
            target_document = Document(streamDoc)
            tmp = len(self.interns_exam_doc)/20
            for i in range(1,tmp+1):
                document = Document(streamDoc)
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if 'tbl_intern' in cell.text:
                                cell.text = cell.text.replace('tbl_intern','tbl_intern%d'%i)
                                break
                # target_document.add_page_break()
                for element in document.element.body:
                    target_document.element.body.append(element)

            tempFile1 = NamedTemporaryFile(delete=False)
            target_document.save(tempFile1.name)
            tempFile1.flush()
            tempFile1.close()

        if docs:
            tpl = None
            if target_document is None:
                stream = BytesIO(docs[0].attachment.decode("base64"))
                tpl = DocxTemplate(stream)
            else:
                tpl = DocxTemplate(tempFile1.name)

            context = {}
            for k in range(0,len(self.interns_exam_doc)/20+1):
                table_interns = []
                for i in range(20*k,20*k+20):
                    if i >= len(self.interns_exam_doc):
                        break
                    intern = self.interns_exam_doc[i]
                    info = {}
                    info['stt'] = str(i+1)
                    info['htk'] = intern.name_without_signal.upper()
                    info['htn'] = intern.name_in_japan
                    if intern.gender == 'nu':
                        info['gt'] = u'女'
                    else:
                        info['gt'] = u'男'
                    info['ns'] = intern_utils.date_time_in_jp(intern.day,intern.month,intern.year)
                    info['t'] = str(intern_utils.get_age_jp(datetime.now(), intern.day,intern.month,intern.year))
                    info['nm'] = intern.blood_group

                    if intern.vision_left:
                        left = 1.5 - (10.0 - intern.vision_left) / 10.0
                        info['tlt'] = "%.1f" % (left)
                    else:
                        info['tlt'] = 'False'
                    if intern.vision_right:
                        right = 1.5 - (10.0 - intern.vision_right) / 10.0
                        info['tlp'] = "%.1f" % (right)
                    else:
                        info['tlp'] = 'False'

                    if intern.preferred_hand == '0':
                        info['tt'] = u'右'
                    elif intern.preferred_hand == '1':
                        info['tt'] = u'左'
                    else:
                        info['tt'] = u'両手'
                    info['cc'] = str(intern.height)
                    info['cn'] = str(intern.weight)
                    info['iq'] = "%s%%"%intern.iq_percentage
                    info['ktk'] = intern.check_kureperin
                    info['hn'] = intern.marital_status.name_in_jp
                    info['pro'] = intern_utils.no_accent_vietnamese(intern.province.name).upper()
                    info['bc'] = intern.certification.name_in_jp
                    table_interns.append(info)

                if k == 0:
                    context['tbl_intern'] = table_interns
                else:
                    context['tbl_intern%d'%k] = table_interns

            context['logo'] = InlineImage(tpl,BytesIO(logo[0].attachment.decode("base64")),width=Mm(35))
            if self.enterprise_doc:
                context['xn'] = self.enterprise_doc.name_jp
            context['nd'] = self.guild.name_in_jp
            context['now'] = intern_utils.date_time_in_jp(datetime.now().day,datetime.now().month,datetime.now().year)

            tpl.render(context)

            tempFile = NamedTemporaryFile(delete=False)
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()

            if tempFile1 is not None:
                os.unlink(tempFile1.name)

            return tempFile
        if tempFile1 is not None:
            os.unlink(tempFile1.name)
        return None



    def set_column_width(self,column, width):
        for cell in column.cells:
            cell.width = width

    def getFieldString(self, model_name, name):
        model_obj = self.env['ir.model']
        model_id = model_obj.search([('model', '=', model_name)])
        if model_id:
            print model_id[0]
            field_ids = self.env['ir.model.fields'].search([('model_id', '=', model_id[0].id), ('name', '=', name)])
            if field_ids:
                return field_ids[0].field_description
        return ""


    # def getFixAllPredefineImage(self,document):
    #     stream = BytesIO(document.attachment.decode("base64"))
    #     doc = Document(stream)
    #     list_image = []
    #     for i, shape in enumerate(doc.inline_shapes):
    #         if shape.type == WD_INLINE_SHAPE.PICTURE:
    #             inline = shape._inline
    #             rId = inline.xpath('./a:graphic/a:graphicData/pic:pic/pic:blipFill/a:blip/@r:embed')[0]
    #             image_part = doc.part.related_parts[rId]
    #             list_image.append(image_part.blob)

    # def calculate_age(self,born):
    #     today = date.today()
    #     return today.year - born.year - ((today.month, today.day) < (born.month, born.day))

    def createCVDoc(self, document, intern, index):
        stream = BytesIO(document.attachment.decode("base64"))
        tpl = DocxTemplate(stream)
        context = {}

        context['ye'] = self.year_expire
        if intern.avatar is not None:
            try:
                streamAvatar = BytesIO(intern.avatar.decode("base64"))
                # context['aa'] = InlineImage(tpl, streamAvatar, width=Mm(20))
                tpl.replace_pic('avatar.jpg',streamAvatar)
            except:
                _logger.info("error")

        context['stt'] = str(index+1)
        # context['ht'] = intern.name
        context['htk'] = intern.name_without_signal.upper()

        if intern.gender == 'nu':
            context['gt'] = '女'
        else:
            context['gt'] = '男'
        context['t'] = str(intern_utils.get_age_jp(datetime.now(),intern.day,intern.month,intern.year))

        if intern.marital_status:
            context['hn'] = intern.marital_status.name_in_jp
        if intern.name_in_japan:
            context['htn'] = intern.name_in_japan
        if intern.date_of_birth:
            context['ns'] = intern_utils.date_time_in_jp(intern.day,intern.month,intern.year)
        if intern.address:
            tmpaddress = intern.address
            if ',' in intern.address:
                tmp = intern.address.split(',')
                tmpaddress = tmp[len(tmp)-1].strip()

            context['dc'] = (intern_utils.no_accent_vietnamese(tmpaddress) + " - " + intern_utils.no_accent_vietnamese(intern.province.name)).upper()
            context['kc'] = intern.province.getDistanceString()
        # if intern.phone_number:
        #     context['sdt'] = intern.phone_number
        if intern.height:
            context['cc'] = str(intern.height)
        if intern.weight:
            context['cn'] = str(intern.weight)
        if intern.vision_left and intern.vision_right:
            left = 1.5-(10.0-intern.vision_left)/10.0
            right = 1.5-(10.0-intern.vision_right)/10.0
            context['tl'] = "%.1f - %.1f"%(left,right)


        if intern.blindness:
            context['cmm'] = CheckedBox().init(24)
            context['kmm'] = CheckBox().init(24)
        else:
            context['cmm'] = CheckBox().init(24)
            context['kmm'] = CheckedBox().init(24)

        if intern.smoking:
            context['cht'] = CheckedBox().init(24)
            context['kht'] = CheckBox().init(24)
        else:
            context['cht'] = CheckBox().init(24)
            context['kht'] = CheckedBox().init(24)

        if intern.preferred_hand == '0':
            context['ptt'] = CheckedBox().init(24)
            context['ttt'] = CheckBox().init(24)
        elif intern.preferred_hand == '1':
            _logger.info('tay phai')
            context['ptt'] = CheckBox().init(24)
            context['ttt'] = CheckedBox().init(24)
        else:
            context['ptt'] = CheckedBox().init(24)
            context['ttt'] = CheckedBox().init(24)

        if intern.surgery:
            context['chx'] = CheckedBox().init(24)
            context['khx'] = CheckBox().init(24)
            context['hxnd'] = intern.surgery_content
        else:
            context['chx'] = CheckBox().init(24)
            context['khx'] = CheckedBox().init(24)

        if intern.drink_alcohol:
            context['crb'] = CheckedBox().init(24)
            context['krb'] = CheckBox().init(24)
        else:
            context['crb'] = CheckBox().init(24)
            context['krb'] = CheckedBox().init(24)


        if intern.certification:
            if intern.certification.id == 1:
                intern.specialized = u'無し'
            elif intern.certification.id == 2:
                intern.specialized = u'無し'
        if intern.specialized:
            context['chn'] = intern_utils.convert_to_docx_string(intern.specialized)
        if intern.favourite:
            context['st'] = intern_utils.convert_to_docx_string(intern.favourite)
        if intern.strong:
            context['dm'] = intern_utils.convert_to_docx_string(intern.strong)
        if intern.weak:
            context['dy'] = intern_utils.convert_to_docx_string(intern.weak)
        if intern.teammate:
            context['kns'] = u'有'
        else:
            context['kns'] = u'無'

        if intern.cooking:
            context['na'] = u'可'
        else:
            context['na'] = u'不可'

        if intern.diseases:
            context['bt'] = u'有'
        else:
            context['bt'] = u'無'
        if intern.blood_group:
            context['nm'] = intern.blood_group
        if intern.check_kureperin:
            context['ktk'] = intern.check_kureperin
        if intern.iq_percentage:
            if u'%' in intern.iq_percentage:
                context['iq'] = intern.iq_percentage
            else:
                context['iq'] = intern.iq_percentage+u"%"
        if intern.family_income:
            context['tng'] = intern.family_income
        if intern.motivation:
            context['dl'] = intern_utils.convert_to_docx_string(intern.motivation)
        if intern.income_after_three_year:
            context['bn']= intern.income_after_three_year
        if intern.job_after_return:
            context['vn'] = intern_utils.convert_to_docx_string(intern.job_after_return)
        if intern.prefer_object:
            context['mg'] = intern_utils.convert_to_docx_string(intern.prefer_object)
        if intern.memory:
            context['kn'] = intern_utils.convert_to_docx_string(intern.memory)
        if intern.valuable:
            context['qg'] = intern_utils.convert_to_docx_string(intern.valuable)

        context['hm'] = CheckBox().init(20)
        context['hs'] = CheckBox().init(20)
        context['lb'] = CheckBox().init(20)
        if intern.education_status:
            if intern.education_status == '1':
                context['hs'] = CheckedBox().init(20)
            elif intern.education_status == '2':
                context['hm'] = CheckedBox().init(20)
            elif intern.education_status == '3':
                context['lb'] = CheckedBox().init(20)
            context['edct'] = intern.education_content

        table_education = []
        for education in sorted(intern.educations,key=lambda x: x.sequence):
            info = {}
            if education.month_start:
                info['nbd'] = intern_utils.date_time_in_jp(month=education.month_start, year=education.year_start)
            else:
                info['nbd'] = intern_utils.date_time_in_jp(year=education.year_start)
            if education.month_end:
                info['nkt'] = intern_utils.date_time_in_jp(month=education.month_end, year=education.year_end)
            else:
                info['nkt'] = intern_utils.date_time_in_jp(year=education.year_end)
            info['tt'] = intern_utils.convert_to_docx_string(intern_utils.no_accent_vietnamese(education.school).upper())
            info['lt'] = education.school_type.name_in_jp
            info['cn'] = intern_utils.convert_to_docx_string(education.specialization)
            info['bc'] = education.certificate.name_in_jp
            if education.graduated:
                info['tn'] = u'卒業'
            else:
                info['tn'] = u'未卒業'

            table_education.append(info)

        context['tbl_educations'] = table_education

        table_employment = []
        if len(intern.employments) == 0:
            table_employment.append({})
            table_employment.append({})
        else:
            for employment in sorted(intern.employments,key=lambda x: x.sequence):
                info = {}
                if employment.month_start:
                    info['nbd'] = intern_utils.date_time_in_jp(month=employment.month_start, year=employment.year_start)
                else:
                    info['nbd'] = intern_utils.date_time_in_jp(year=employment.year_start)
                if employment.month_end:
                    info['nkt'] = intern_utils.date_time_in_jp(month=employment.month_end, year=employment.year_end)
                else:
                    info['nkt'] = intern_utils.date_time_in_jp(year=employment.year_end)
                info['ct'] = intern_utils.convert_to_docx_string(employment.company)
                info['cv'] = intern_utils.convert_to_docx_string(employment.description)
                table_employment.append(info)

        context['tbl_employ'] = table_employment

        table_family = []
        for i,person in enumerate(sorted(intern.family_members,key=lambda x: x.sequence)):
            if i < 5:
                context['p%dht'%(i+1)] = intern_utils.no_accent_vietnamese(person.name).upper()
                context['p%dqh'%(i+1)] = person.relationship
                if person.birth_year>0:
                    context['p%ddt'%(i+1)] = str((datetime.now().year) - int(person.birth_year))
                # context['p%ddt' % (i + 1)] = person.ages
                if person.job:
                    context['p%dnn' % (i + 1)] = intern_utils.convert_to_docx_string(person.job)
                if person.live_together:
                    context['p%dsc' % (i + 1)] = Tick()
                else:
                    context['p%dsr' % (i + 1)] = Tick()
            else:

                if person.live_together:
                    job = person.job
                    if not person.job:
                        job=""
                    table_family.append({'ht': intern_utils.no_accent_vietnamese(person.name).upper(), 'qh': person.relationship,
                                         'dt': str((datetime.now().year)- int(person.birth_year)),
                                         # 'dt': person.ages,
                                             'nn': job, 'sc': Tick(),'sr':''})
                else:
                    job = person.job
                    if not person.job:
                        job = ""
                    table_family.append({'ht': intern_utils.no_accent_vietnamese(person.name).upper(), 'qh': person.relationship,
                                         'dt': str((datetime.now().year)  - int(person.birth_year)),
                                         # 'dt': person.ages,
                                         'nn': job, 'sc': '', 'sr': Tick()})
        context['tbl_family'] = table_family

        if intern.family_member_in_jp:
            context['cnt'] = CheckedBox().init(24)
            context['knt'] = CheckBox().init(24)
            context['nton'] = intern.family_member_in_jp
        else:
            context['cnt'] = CheckBox().init(24)
            context['knt'] = CheckedBox().init(24)

        if intern.family_accept:
            context['gdy'] = CheckedBox().init(24)
            context['gpd'] = CheckBox().init(24)
        else:
            context['gdy'] = CheckBox().init(24)
            context['gpd'] = CheckedBox().init(24)
        tpl.set_fix_table_border()
        tpl.render(context)

        # byteIoRespond = BytesIO()
        # tpl.save(byteIoRespond)
        # byteIoRespond.seek(0)
        # return byteIoRespond
        tempFile = NamedTemporaryFile(delete=False)
        tpl.save(tempFile)
        tempFile.flush()
        tempFile.close()
        return tempFile

    @api.model
    def on_order_changed(self,id, column, state):
        invoice = self.env['intern.invoice'].browse(int(id))
        if invoice:
            if state==0:
                invoice.write({'order': "%s"%(column)})
            else:
                invoice.write({'order': "%s desc" % (column)})



    job = fields.Char("Ngành nghề")
    year_expire = fields.Integer("Thời hạn hợp đồng (năm)")
    salary_base = fields.Char("Lương cơ bản")
    salary_real = fields.Char("Lương thực lĩnh")
    subsidize_start_month = fields.Integer("Trợ cấp đầu tháng")
    number_man = fields.Integer("Số lượng nam",default=0)
    number_women = fields.Integer("Số lượng nữ",default=0)
    number_total = fields.Integer("Số lượng trúng tuyển",default=0)

    @api.multi
    @api.onchange('number_man','number_women')
    def _compute_number_total(self):
        for rec in self:
            rec.number_total = rec.number_man +rec.number_women

    source_man = fields.Integer("Nguồn nam",default=0)
    source_women = fields.Integer("Nguồn nữ",default=0)
    source_total = fields.Integer("Số lượng thi tuyển",default=0)

    @api.multi
    @api.onchange('source_man', 'source_women')
    def _compute_source_total(self):
        for rec in self:
            rec.source_total = rec.source_man + rec.source_women

    age_from = fields.Integer("Tuổi từ")
    age_to = fields.Integer("Đến tuổi")
    certificate = fields.Many2one("intern.certification","Trình độ")
    marital = fields.Many2one("marital","Hôn nhân")
    vision = fields.Char("Thị lực")
    physical = fields.Char("Thể lực")
    preferred_hand = fields.Selection([('0', 'Tay phải'), ('1', 'Tay trái')], string="Tay thuận", default='0')
    height_man = fields.Integer("Chiều cao nam")
    height_women = fields.Integer("Chiều cao nữ")
    weight_man = fields.Integer("Cân nặng nam")
    weight_women = fields.Integer("Cân nặng nữ")
    smoking = fields.Boolean("Hút thuốc")

    job_description = fields.Char("Nội dung công việc")
    other_requirement = fields.Char("Tiêu chuẩn khác")

    note = fields.Char("Ghi chú")

    @api.model
    def _get_current_year(self):
        return str(datetime.now().year)

    #Thời hạn
    day_supply_form = fields.Char("Ngày", size=2)
    month_supply_form = fields.Selection([('01', '01'), ('02', '02'), ('03', '03'), ('04', '04'),
                                        ('05', '05'), ('06', '06'), ('07', '07'), ('08', '08'),
                                        ('09', '09'), ('10', '10'), ('11', '11'), ('12', '12'), ], "Tháng")

    year_supply_form = fields.Char("Năm", size=4, default=lambda self: self._get_current_year())

    date_supply_form = fields.Char("Ngày cấp form", store=False, compute='_date_supply_form')

    @api.one
    @api.depends('day_supply_form', 'month_supply_form', 'year_supply_form')
    def _date_supply_form(self):
        if self.day_supply_form and self.month_supply_form and self.year_supply_form:
            self.date_supply_form = u"Ngày %s tháng %s năm %s" % (self.day_supply_form, self.month_supply_form, self.year_supply_form)
        elif self.month_supply_form and self.year_supply_form:
            self.date_supply_form = u"Tháng %s năm %s" % (
                                                 self.month_supply_form, self.year_supply_form)
        elif self.year_supply_form:
            self.date_supply_form = u'Năm %s'%self.year_supply_form
        else:
            self.date_supply_form = ""

    #
    day_check_form = fields.Char("Ngày", size=2)
    month_check_form = fields.Selection([('01', '01'), ('02', '02'), ('03', '03'), ('04', '04'),
                                        ('05', '05'), ('06', '06'), ('07', '07'), ('08', '08'),
                                        ('09', '09'), ('10', '10'), ('11', '11'), ('12', '12'), ], "Tháng")

    year_check_form = fields.Char("Năm", size=4, default=lambda self: self._get_current_year())

    date_check_form = fields.Char("Ngày check form", store=False, compute='_date_check_form')

    @api.one
    @api.depends('day_check_form', 'month_check_form', 'year_check_form')
    def _date_check_form(self):
        if self.day_check_form and self.month_check_form and self.year_check_form:
            self.date_check_form = u"Ngày %s tháng %s năm %s" % (
                self.day_check_form, self.month_check_form, self.year_check_form)
        elif self.month_check_form and self.year_check_form:
            self.date_check_form = u"Tháng %s năm %s" % (
                self.month_check_form, self.year_check_form)
        elif self.year_check_form:
            self.date_check_form = u'Năm %s' % self.year_check_form
        else:
            self.date_check_form = ""

    #
    day_send_form = fields.Char("Ngày", size=2)
    month_send_form = fields.Selection([('01', '01'), ('02', '02'), ('03', '03'), ('04', '04'),
                                        ('05', '05'), ('06', '06'), ('07', '07'), ('08', '08'),
                                        ('09', '09'), ('10', '10'), ('11', '11'), ('12', '12'), ], "Tháng")

    year_send_form = fields.Char("Năm", size=4, default=lambda self: self._get_current_year())

    date_send_form = fields.Char("Ngày gửi form", store=False, compute='_date_send_form')

    @api.one
    @api.depends('day_send_form', 'month_send_form', 'year_send_form')
    def _date_send_form(self):
        if self.day_send_form and self.month_send_form and self.year_send_form:
            self.date_send_form = u"Ngày %s tháng %s năm %s" % (
                self.day_send_form, self.month_send_form, self.year_send_form)
        elif self.month_send_form and self.year_send_form:
            self.date_send_form = u"Tháng %s năm %s" % (
                self.month_send_form, self.year_send_form)
        elif self.year_check_form:
            self.date_send_form = u'Năm %s' % self.year_send_form
        else:
            self.date_send_form = ""

    #
    day_exam = fields.Char("Ngày", size=2)
    month_exam = fields.Selection([('01', '01'), ('02', '02'), ('03', '03'), ('04', '04'),
                                        ('05', '05'), ('06', '06'), ('07', '07'), ('08', '08'),
                                        ('09', '09'), ('10', '10'), ('11', '11'), ('12', '12'), ], "Tháng")

    year_exam = fields.Char("Năm", size=4, default=lambda self: self._get_current_year())

    date_exam = fields.Char("Ngày thi", store=False, compute='_date_exam')

    @api.one
    @api.depends('day_exam', 'month_exam', 'year_exam')
    def _date_exam(self):
        if self.day_exam and self.month_exam and self.year_exam:
            self.date_exam = u"Ngày %s tháng %s năm %s" % (
                self.day_exam, self.month_exam, self.year_exam)
        elif self.month_exam and self.year_exam:
            self.date_exam = u"Tháng %s năm %s" % (
                self.month_exam, self.year_exam)
        elif self.year_exam:
            self.date_exam = u'Năm %s' % self.year_exam
        else:
            self.date_exam = ""

    #
    date_exam_short = fields.Date("Ngày thi tuyển", store=True, compute='_date_exam_short')

    @api.multi
    @api.depends('day_exam', 'month_exam', 'year_exam')
    def _date_exam_short(self):
        for rec in self:
            if rec.day_exam and rec.month_exam and rec.year_exam:
                rec.date_exam_short = datetime.strptime('%s-%s-%s' % (rec.year_exam, rec.month_exam, rec.day_exam),
                                                        '%Y-%m-%d')
            else:
                rec.date_exam_short = None

    # day_departure = fields.Char("Ngày", size=2)
    # month_departure = fields.Selection([('01', '01'), ('02', '02'), ('03', '03'), ('04', '04'),
    #                                     ('05', '05'), ('06', '06'), ('07', '07'), ('08', '08'),
    #                                     ('09', '09'), ('10', '10'), ('11', '11'), ('12', '12'), ], "Tháng")
    #
    # year_departure = fields.Char("Năm", size=4, default=lambda self: self._get_current_year())
    #
    # date_departure = fields.Char("Ngày xuất cảnh Dự kiến", store=False, compute='_date_departure')
    date_departure = fields.Date("Ngày xuất cảnh Dự kiến")

    # @api.one
    # @api.depends('day_departure', 'month_departure', 'year_departure')
    # def _date_departure(self):
    #     if self.day_departure and self.month_departure and self.year_departure:
    #         self.date_departure = u"Ngày %s tháng %s năm %s" % (
    #             self.day_departure, self.month_departure, self.year_departure)
    #     elif self.month_departure and self.year_departure:
    #         self.date_departure = u"Tháng %s năm %s" % (
    #             self.month_departure, self.year_departure)
    #     elif self.year_departure:
    #         self.date_departure = u'Năm %s' % self.year_departure
    #     else:
    #         self.date_departure = ""

    room_responsive = fields.Char("Phòng NB phụ trách")

    day_finish = fields.Char("Ngày", size=2)
    month_finish = fields.Selection([('01', '01'), ('02', '02'), ('03', '03'), ('04', '04'),
                                        ('05', '05'), ('06', '06'), ('07', '07'), ('08', '08'),
                                        ('09', '09'), ('10', '10'), ('11', '11'), ('12', '12'), ], "Tháng")

    year_finish = fields.Char("Năm", size=4, default=lambda self: self._get_current_year())

    date_finish = fields.Char("Ngày hoàn thành", store=False, compute='_date_finish')

    @api.one
    @api.depends('day_finish', 'month_finish', 'year_finish')
    def _date_finish(self):
        if self.day_finish and self.month_finish and self.year_finish:
            self.date_finish = u"Ngày %s tháng %s năm %s" % (
                self.day_finish, self.month_finish, self.year_finish)
        elif self.month_finish and self.year_finish:
            self.date_finish = u"Tháng %s năm %s" % (
                self.month_finish, self.year_finish)
        elif self.year_finish:
            self.date_finish = u'Năm %s' % self.year_finish
        else:
            self.date_finish = ""

    #
    day_cancel_tclt = fields.Char("Ngày", size=2)
    month_cancel_tclt = fields.Selection([('01', '01'), ('02', '02'), ('03', '03'), ('04', '04'),
                                        ('05', '05'), ('06', '06'), ('07', '07'), ('08', '08'),
                                        ('09', '09'), ('10', '10'), ('11', '11'), ('12', '12'), ], "Tháng")

    year_cancel_tclt = fields.Char("Năm", size=4, default=lambda self: self._get_current_year())

    date_cancel_tclt = fields.Char("Ngày hủy TCLT", store=False, compute='_date_cancel_tclt')

    @api.one
    @api.depends('day_cancel_tclt', 'month_cancel_tclt', 'year_cancel_tclt')
    def _date_cancel_tclt(self):
        if self.day_cancel_tclt and self.month_cancel_tclt and self.year_cancel_tclt:
            self.date_cancel_tclt = u"Ngày %s tháng %s năm %s" % (
                self.day_cancel_tclt, self.month_cancel_tclt, self.year_cancel_tclt)
        elif self.month_cancel_tclt and self.year_departure:
            self.date_cancel_tclt = u"Tháng %s năm %s" % (
                self.month_cancel_tclt, self.year_cancel_tclt)
        elif self.year_cancel_tclt:
            self.date_cancel_tclt = u'Năm %s' % self.year_cancel_tclt
        else:
            self.date_cancel_tclt = ""

    #
    day_cancel_visa_xc = fields.Char("Ngày", size=2)
    month_cancel_visa_xc = fields.Selection([('01', '01'), ('02', '02'), ('03', '03'), ('04', '04'),
                                        ('05', '05'), ('06', '06'), ('07', '07'), ('08', '08'),
                                        ('09', '09'), ('10', '10'), ('11', '11'), ('12', '12'), ], "Tháng")

    year_cancel_visa_xc = fields.Char("Năm", size=4, default=lambda self: self._get_current_year())

    date_cancel_visa_xc = fields.Char("Ngày hủy Visa-XC", store=False, compute='_date_cancel_visa_xc')

    @api.one
    @api.depends('day_cancel_visa_xc', 'month_cancel_visa_xc', 'year_cancel_visa_xc')
    def _date_cancel_visa_xc(self):
        if self.day_cancel_visa_xc and self.month_cancel_visa_xc and self.year_cancel_visa_xc:
            self.date_cancel_visa_xc = u"Ngày %s tháng %s năm %s" % (
                self.day_cancel_visa_xc, self.month_cancel_visa_xc, self.year_cancel_visa_xc)
        elif self.month_cancel_visa_xc and self.year_cancel_visa_xc:
            self.date_cancel_visa_xc = u"Tháng %s năm %s" % (
                self.month_cancel_visa_xc, self.year_cancel_visa_xc)
        elif self.year_cancel_tclt:
            self.date_cancel_visa_xc = u'Năm %s' % self.year_cancel_visa_xc
        else:
            self.date_cancel_visa_xc = ""


    #
    day_departure2 = fields.Char("Ngày", size=2)
    month_departure2 = fields.Selection([('01', '01'), ('02', '02'), ('03', '03'), ('04', '04'),
                                        ('05', '05'), ('06', '06'), ('07', '07'), ('08', '08'),
                                        ('09', '09'), ('10', '10'), ('11', '11'), ('12', '12'), ], "Tháng")

    year_departure2 = fields.Char("Năm", size=4, default=lambda self: self._get_current_year())

    date_departure2 = fields.Char("Ngày xuất cảnh", store=False, compute='_date_departure2')

    @api.one
    @api.depends('day_departure2', 'month_departure2', 'year_departure2')
    def _date_departure2(self):
        if self.day_departure2 and self.month_departure2 and self.year_departure2:
            self.date_departure2 = u"Ngày %s tháng %s năm %s" % (
                self.day_departure2, self.month_departure2, self.year_departure2)
        elif self.month_departure2 and self.year_departure2:
            self.date_departure2 = u"Tháng %s năm %s" % (
                self.month_departure2, self.year_departure2)
        elif self.year_departure2:
            self.date_departure2 = u'Năm %s' % self.year_departure2
        else:
            self.date_departure2 = ""


    reason_cancel_tclt_visa_xc =fields.Char("Lý do đơn hàng bị hủy TCLT - Visa - XC")


    # create_date_only = fields.Char("Ngày tạo",store=False,compute='_get_create_date')

    # @api.one
    # def _get_create_date(self):
    #     if self.create_date is not None:
    #         create_date_only = self.create_date[10:]





    #Thong tin bo sung cho ho so noi

    day_create_letter_promotion = fields.Char("Ngày", size=2)
    month_create_letter_promotion = fields.Selection([('01', '01'), ('02', '02'), ('03', '03'), ('04', '04'),
                                        ('05', '05'), ('06', '06'), ('07', '07'), ('08', '08'),
                                        ('09', '09'), ('10', '10'), ('11', '11'), ('12', '12'), ], "Tháng")

    year_create_letter_promotion = fields.Char("Năm", size=4, default=lambda self: self._get_current_year())
    date_create_letter_promotion = fields.Char("Ngày làm thư tiến cử", store=False, compute='_date_create_letter_pro')

    @api.multi
    @api.depends('day_create_letter_promotion', 'month_create_letter_promotion', 'year_create_letter_promotion')
    def _date_create_letter_pro(self):

        for rec in self:
            if rec.day_create_letter_promotion and rec.month_create_letter_promotion and rec.year_create_letter_promotion:
                rec.date_create_letter_promotion = u"Ngày %s tháng %s năm %s" % (
                    rec.day_create_letter_promotion, rec.month_create_letter_promotion, rec.year_create_letter_promotion)
            elif rec.month_create_letter_promotion and rec.year_create_letter_promotion:
                rec.date_create_letter_promotion = u"Tháng %s năm %s" % (
                    rec.month_create_letter_promotion, rec.year_create_letter_promotion)
            elif rec.year_create_letter_promotion:
                rec.date_create_letter_promotion = u'Năm %s' % rec.year_create_letter_promotion
            else:
                rec.date_create_letter_promotion = ""



    day_create_plan_training = fields.Char("Ngày", size=2)
    month_create_plan_training = fields.Selection([('01', '01'), ('02', '02'), ('03', '03'), ('04', '04'),
                                        ('05', '05'), ('06', '06'), ('07', '07'), ('08', '08'),
                                        ('09', '09'), ('10', '10'), ('11', '11'), ('12', '12'), ], "Tháng")

    year_create_plan_training = fields.Char("Năm", size=4, default=lambda self: self._get_current_year())
    date_create_plan_training = fields.Char("Ngày lập kế hoạch đào tạo theo ủy thác(Trước ngày bắt đầu ít nhất 1 ngày, chú ý Chủ Nhật)", store=False, compute='_date_create_plan_training')


    @api.multi
    @api.depends('day_create_plan_training', 'month_create_plan_training', 'year_create_plan_training')
    def _date_create_plan_training(self):
        for rec in self:
            if rec.day_create_plan_training and rec.month_create_plan_training and rec.year_create_plan_training:
                rec.date_create_plan_training = u"Ngày %s tháng %s năm %s" % (
                    rec.day_create_plan_training, rec.month_create_plan_training, rec.year_create_plan_training)

            elif rec.month_create_plan_training and rec.year_create_plan_training:
                rec.date_create_plan_training = u"Tháng %s năm %s" % (
                    rec.month_create_plan_training, rec.year_create_plan_training)
            elif rec.year_create_plan_training:
                rec.date_create_plan_training = u'Năm %s' % rec.year_create_plan_training
            else:
                rec.date_create_plan_training = ""

    day_start_training = fields.Char("Ngày", size=2)
    month_start_training = fields.Selection([('01', '01'), ('02', '02'), ('03', '03'), ('04', '04'),
                                        ('05', '05'), ('06', '06'), ('07', '07'), ('08', '08'),
                                        ('09', '09'), ('10', '10'), ('11', '11'), ('12', '12'), ], "Tháng")

    year_start_training = fields.Char("Năm", size=4, default=lambda self: self._get_current_year())

    date_start_training = fields.Char("Ngày bắt đầu khóa học chú ý Chủ Nhật(Trước ngày bắt đầu 1 tháng)", store=False, compute='_date_start_training')


    @api.multi
    @api.depends('day_start_training', 'month_start_training', 'year_start_training')
    def _date_start_training(self):
        for rec in self:
            if rec.day_start_training and rec.month_start_training and rec.year_start_training:
                rec.date_start_training = u"Ngày %s tháng %s năm %s" % (
                    rec.day_start_training, rec.month_start_training, rec.year_start_training)

                tmp_date = datetime.strptime('%s/%s/%s' % (
                rec.day_start_training, rec.month_start_training, rec.year_start_training),
                                             '%d/%m/%Y')

                tmp_date_create_plan = tmp_date - relativedelta(days=3)
                # if tmp_date_create_plan.day<10:
                #     rec.day_create_plan_training = '0%d'%tmp_date_create_plan.day
                # else:
                if not rec.day_create_plan_training or not rec.month_create_plan_training or not rec.year_create_plan_training:
                    rec.day_create_plan_training = '%02d'%tmp_date_create_plan.day
                    rec.month_create_plan_training = '%02d'%tmp_date_create_plan.month
                    rec.year_create_plan_training = '%d'%tmp_date_create_plan.year

                tmp_date_end_plan = tmp_date + relativedelta(months=1)

                # if tmp_date_end_plan.day<10:
                #     rec.day_end_training = '0%d'%tmp_date_end_plan.day
                # else:
                if not rec.day_end_training or not rec.month_end_training or not rec.year_end_training:
                    rec.day_end_training = '%02d'%tmp_date_end_plan.day

                    # if tmp_date_end_plan.month < 10:
                    #     rec.month_end_training = '0%d'%tmp_date_end_plan.month
                    # else:
                    rec.month_end_training = '%02d'%tmp_date_end_plan.month
                    rec.year_end_training = '%d'%tmp_date_end_plan.year

                tmp_date_report_customer = tmp_date_end_plan + relativedelta(days=1)

                if not rec.day_create_plan_training_report_customer or not rec.month_create_plan_training_report_customer or not rec.year_create_plan_training_report_customer:
                    rec.day_create_plan_training_report_customer = '%02d'%tmp_date_report_customer.day
                    # else:
                    #     rec.day_create_plan_training_report_customer = '%d'%tmp_date_report_customer.day

                    # if tmp_date_report_customer.month < 10:
                    rec.month_create_plan_training_report_customer = '%02d'%tmp_date_report_customer.month
                    # else:
                    #     rec.month_create_plan_training_report_customer = '%d'%tmp_date_report_customer.month
                    rec.year_create_plan_training_report_customer = '%d'%tmp_date_report_customer.year

                # day_create_plan_training_report_customer
                if not rec.day_pay_finance1 or not rec.month_pay_finance1 or not rec.year_pay_finance1:
                    rec.day_pay_finance1 = rec.day_create_plan_training
                    rec.month_pay_finance1 = rec.month_create_plan_training
                    rec.year_pay_finance1 = rec.year_create_plan_training

            elif rec.month_start_training and rec.year_start_training:
                rec.date_start_training = u"Tháng %s năm %s" % (
                    rec.month_start_training, rec.year_start_training)
            elif rec.year_start_training:
                rec.date_start_training = u'Năm %s' % rec.year_start_training
            else:
                rec.date_start_training = ""

    day_end_training = fields.Char("Ngày", size=2)
    month_end_training = fields.Selection([('01', '01'), ('02', '02'), ('03', '03'), ('04', '04'),
                                        ('05', '05'), ('06', '06'), ('07', '07'), ('08', '08'),
                                        ('09', '09'), ('10', '10'), ('11', '11'), ('12', '12'), ], "Tháng")

    year_end_training = fields.Char("Năm", size=4, default=lambda self: self._get_current_year())
    date_end_training = fields.Char("Ngày kết thúc khóa học(chú ý Chủ Nhật) Tính từ thời điểm làm hồ sơ sau khoảng 10 ngày", store=False, compute='_date_end_training')



    @api.multi
    @api.depends('day_end_training', 'month_end_training', 'year_end_training')
    def _date_end_training(self):
        for rec in self:
            if rec.day_end_training and rec.month_end_training and rec.year_end_training:
                rec.date_end_training = u"Ngày %s tháng %s năm %s" % (
                    rec.day_end_training, rec.month_end_training, rec.year_end_training)
            elif rec.month_end_training and rec.year_end_training:
                rec.date_end_training = u"Tháng %s năm %s" % (
                    rec.month_end_training, rec.year_end_training)
            elif rec.year_end_training:
                rec.date_end_training = u'Năm %s' % rec.year_end_training
            else:
                rec.date_end_training = ""

    day_create_plan_training_report_customer = fields.Char("Ngày", size=2)
    month_create_plan_training_report_customer = fields.Selection([('01', '01'), ('02', '02'), ('03', '03'), ('04', '04'),
                                        ('05', '05'), ('06', '06'), ('07', '07'), ('08', '08'),
                                        ('09', '09'), ('10', '10'), ('11', '11'), ('12', '12'), ], "Tháng")

    year_create_plan_training_report_customer = fields.Char("Năm", size=4, default=lambda self: self._get_current_year())
    date_create_plan_training_report_customer = fields.Char("Ngày lập kế hoạch đào tạo BÁO CÁO KHÁCH HÀNG trước ngày bắt đầu ít nhất 1 ngày chú ý CHỦ NHẬT", store=False, compute='_date_create_plan_training_report_customer')


    @api.multi
    @api.depends('day_create_plan_training_report_customer', 'month_create_plan_training_report_customer', 'year_create_plan_training_report_customer')
    def _date_create_plan_training_report_customer(self):
        for rec in self:
            if rec.day_create_plan_training_report_customer and rec.month_create_plan_training_report_customer and rec.year_create_plan_training_report_customer:
                rec.date_create_plan_training_report_customer = u"Ngày %s tháng %s năm %s" % (
                    rec.day_create_plan_training_report_customer, rec.month_create_plan_training_report_customer, rec.year_create_plan_training_report_customer)
            elif rec.month_create_plan_training_report_customer and rec.year_create_plan_training_report_customer:
                rec.date_create_plan_training_report_customer = u"Tháng %s năm %s" % (
                    rec.month_create_plan_training_report_customer, rec.year_create_plan_training_report_customer)
            elif rec.year_create_plan_training_report_customer:
                rec.date_create_plan_training_report_customer = u'Năm %s' % rec.year_create_plan_training_report_customer
            else:
                rec.date_create_plan_training_report_customer = ""

    day_pay_finance1 = fields.Char("Ngày", size=2)
    month_pay_finance1 = fields.Selection([('01', '01'), ('02', '02'), ('03', '03'), ('04', '04'),
                                        ('05', '05'), ('06', '06'), ('07', '07'), ('08', '08'),
                                        ('09', '09'), ('10', '10'), ('11', '11'), ('12', '12'), ], "Tháng")

    year_pay_finance1 = fields.Char("Năm", size=4, default=lambda self: self._get_current_year())
    date_pay_finance1 = fields.Char("Ngày nộp tài chính đợt 1 (Sau ngày trúng tuyển 7 ngày)", store=False, compute='_date_pay_finance1')



    @api.multi
    @api.depends('day_pay_finance1', 'month_pay_finance1',
                 'year_pay_finance1')
    def _date_pay_finance1(self):
        for rec in self:
            if rec.day_pay_finance1 and rec.month_pay_finance1 and rec.year_pay_finance1:
                rec.date_pay_finance1 = u"Ngày %s tháng %s năm %s" % (
                    rec.day_pay_finance1, rec.month_pay_finance1,
                    rec.year_pay_finance1)
            elif rec.month_pay_finance1 and rec.year_pay_finance1:
                rec.date_pay_finance1 = u"Tháng %s năm %s" % (
                    rec.month_pay_finance1, rec.year_pay_finance1)
            elif rec.year_pay_finance1:
                rec.date_pay_finance1 = u'Năm %s' % rec.year_pay_finance1
            else:
                rec.date_pay_finance1 = ""

    day_pay_finance2 = fields.Char("Ngày", size=2)
    month_pay_finance2 = fields.Selection([('01', '01'), ('02', '02'), ('03', '03'), ('04', '04'),
                                        ('05', '05'), ('06', '06'), ('07', '07'), ('08', '08'),
                                        ('09', '09'), ('10', '10'), ('11', '11'), ('12', '12'), ], "Tháng")

    year_pay_finance2 = fields.Char("Năm", size=4, default=lambda self: self._get_current_year())
    date_pay_finance2 = fields.Char("Ngày nộp tài chính đợt 2 (Trước ngày dự kiến xuất cảnh 10 ngày)", store=False, compute='_date_pay_finance2')



    @api.multi
    @api.depends('day_pay_finance2', 'month_pay_finance2',
                 'year_pay_finance2')
    def _date_pay_finance2(self):
        for rec in self:
            if rec.day_pay_finance2 and rec.month_pay_finance2 and rec.year_pay_finance2:
                rec.date_pay_finance2 = u"Ngày %s tháng %s năm %s" % (
                    rec.day_pay_finance2, rec.month_pay_finance2,
                    rec.year_pay_finance2)
            elif rec.month_pay_finance2 and rec.year_pay_finance2:
                rec.date_pay_finance2 = u"Tháng %s năm %s" % (
                    rec.month_pay_finance2, rec.year_pay_finance2)
            elif rec.year_pay_finance2:
                rec.date_pay_finance2 = u'Năm %s' % rec.year_pay_finance2
            else:
                rec.date_pay_finance2 = ""



    length_training = fields.Char("Thời gian đào tạo khóa học (Chữ Hán)")
    hours_training = fields.Integer("Tổng số thời gian đào tạo (số giờ học)")

    training_center = fields.Many2one('trainingcenter',string='Trung tâm đào tạo')
    guild = fields.Many2one('intern.guild',string='Nghiệp đoàn')

    enterprise_doc = fields.Many2one('intern.enterprise',string='Xí nghiệp')

    @api.multi
    @api.onchange('enterprise_doc')
    def _onchange_enterprise_doc(self):
        for rec in self:
            for intern in rec.interns_clone:
                intern.enterprise = rec.enterprise_doc

    # @api.multi
    # @api.onchange('interns_clone')
    # def _onchange_interns_clone(self):
    #     for rec in self:
    #         if rec.enterprise_doc:
    #             for intern in rec.interns_clone:
    #                 intern.enterprise = rec.enterprise_doc


    dispatchcom2 = fields.Many2many('dispatchcom2',string=u'Công ty phái cử thứ 2')

    @api.multi
    @api.onchange('dispatchcom2')
    # @api.depends('dispatchcom2')
    def _onchange_dispatchcom2(self):
        for rec in self:
            for intern in rec.interns_clone:
                if rec.dispatchcom2:
                    # intern.dispatchcom2 = (6, 0, [rec.dispatchcom2[0].id])
                    intern.dispatchcom2 = rec.dispatchcom2[0]
                else:
                    intern.dispatchcom2 = False
            for intern in rec.interns_pass_doc:
                if rec.dispatchcom2:
                    intern.dispatchcom2 = rec.dispatchcom2[0]
                else:
                    intern.dispatchcom2 = False
            for intern in rec.interns_pass_doc_hs:
                if rec.dispatchcom2:
                    intern.dispatchcom2 = rec.dispatchcom2[0]
                else:
                    intern.dispatchcom2 = False

    dispatchcom1 = fields.Many2one('dispatchcom1','Pháp nhân')



    name_working_department = fields.Char("Tên bộ phận TTS sẽ làm việc trong xí nghiệp (có trong hợp đồng lương) - Tiếng Nhật ")

    job_predefine = fields.Many2one('intern.job','Ngành nghề xin thư tiến cử')



    job_en = fields.Char("Ngành nghề xin thư tiến cử lấy từ hợp đồng lương (Tiếng Anh)")
    job_jp = fields.Char("Ngành nghề xin thư tiến cử lấy từ hợp đồng lương (Tiếng Nhật)")
    job_vi = fields.Char("Ngành nghề xin thư tiến cử lấy từ hợp đồng lương (Tiếng Việt)")

    @api.onchange('job_predefine')
    def job_change(self):
        if self.job_predefine:
            self.job_en = self.job_predefine.name_en
            self.job_jp = self.job_predefine.name_jp
            self.job_vi = self.job_predefine.name


    day_departure_doc = fields.Char("Ngày", size=2)
    month_departure_doc = fields.Selection([('01', '01'), ('02', '02'), ('03', '03'), ('04', '04'),
                                        ('05', '05'), ('06', '06'), ('07', '07'), ('08', '08'),
                                        ('09', '09'), ('10', '10'), ('11', '11'), ('12', '12'), ], "Tháng")

    year_departure_doc = fields.Char("Năm", size=4, default=lambda self: self._get_current_year())

    date_departure_doc = fields.Char("Ngày xuất cảnh Dự kiến", store=False, compute='_date_departure_doc')


    @api.multi
    @api.depends('day_departure_doc', 'month_departure_doc', 'year_departure_doc')
    def _date_departure_doc(self):
        for rec in self:
            if rec.day_departure_doc and rec.month_departure_doc and rec.year_departure_doc:
                rec.date_departure_doc = u"Ngày %s tháng %s năm %s" % (
                    rec.day_departure_doc, rec.month_departure_doc, rec.year_departure_doc)
            elif rec.month_departure_doc and rec.year_departure_doc:
                rec.date_departure_doc = u"Tháng %s năm %s" % (
                    rec.month_departure_doc, rec.year_departure_doc)
            elif rec.year_departure_doc:
                rec.date_departure_doc = u'Năm %s' % rec.year_departure_doc
            else:
                rec.date_departure_doc = ""

    person_sign_proletter = fields.Char("Tên người ký thư tiến cử (kiểm tra thường xuyên tránh sai)",default=u"Vũ Trường Giang")
    position_person_sign = fields.Char("Chức danh người ký thư tiến cử - Tiếng Anh",default=u'Head of Division for Japan -  Southeast Asia')
    position_person_sign_jp= fields.Char("Chức danh người ký thư tiến cử - Tiếng Nhật",default=u'日本東南アジア部長')
    developing_employee = fields.Char("Cán bộ PTTT")

    back_to_pc2 = fields.Boolean('Sau khi về nước sẽ quay lại Cty PC2 ')

    day_sign_proletter = fields.Char("Ngày", size=2)
    month_sign_proletter = fields.Selection([('01', '01'), ('02', '02'), ('03', '03'), ('04', '04'),
                                           ('05', '05'), ('06', '06'), ('07', '07'), ('08', '08'),
                                           ('09', '09'), ('10', '10'), ('11', '11'), ('12', '12'), ], "Tháng")

    year_sign_proletter = fields.Char("Năm", size=4, default=lambda self: self._get_current_year())
    date_sign_proletter = fields.Char("Ngày ký hợp đồng phái cử", store=False,
                                    compute='_date_sign_proletter')     #a122

    @api.multi
    @api.depends('day_sign_proletter', 'month_sign_proletter',
                 'year_sign_proletter')
    def _date_sign_proletter(self):
        for rec in self:
            if rec.day_sign_proletter and rec.month_sign_proletter and rec.year_sign_proletter:
                rec.date_sign_proletter = u"Ngày %s tháng %s năm %s" % (
                    rec.day_sign_proletter, rec.month_sign_proletter,
                    rec.year_sign_proletter)
            elif rec.month_sign_proletter and rec.year_sign_proletter:
                rec.date_sign_proletter = u"Tháng %s năm %s" % (
                    rec.month_sign_proletter, rec.year_sign_proletter)
            elif rec.year_pay_finance2:
                rec.date_sign_proletter = u'Năm %s' % rec.year_sign_proletter
            else:
                rec.date_sign_proletter = ""

    date_create_letter_promotion_short = fields.Date('Ngày làm thư tiến cử', store=True,
                                                     compute='_date_create_letter_promotion_short')

    @api.multi
    @api.depends('day_create_letter_promotion', 'month_create_letter_promotion', 'year_create_letter_promotion')
    def _date_create_letter_promotion_short(self):
        for rec in self:
            if rec.day_create_letter_promotion and rec.month_create_letter_promotion and rec.year_create_letter_promotion:
                rec.date_create_letter_promotion_short = datetime.strptime('%s-%s-%s' %
                                                                           (rec.year_create_letter_promotion,
                                                                            rec.month_create_letter_promotion,
                                                                            rec.day_create_letter_promotion),
                                                                           '%Y-%m-%d')
            else:
                rec.date_create_letter_promotion_short = None


    def create_doc_1_3(self, intern, index):
        docs = self.env['intern.document'].search([('name', '=', "Doc1-3")], limit=1)
        if docs:
            stream = BytesIO(docs[0].attachment.decode("base64"))
            tpl = DocxTemplate(stream)
            context = {}

             #ovan 3 la Nu, 1 la nam, 2 la ko, 5 la co

            if intern.gender == 'nam':
                tpl.remove_shape(u'id="1" name="Oval 1"')

            else:
                tpl.remove_shape(u'id="5" name="Oval 5"')

            if intern.marital_status.id is not 2:
                tpl.remove_shape(u'id="2" name="Oval 2"')
            else:
                tpl.remove_shape(u'id="6" name="Oval 6"')

            context['a3'] = intern_utils.no_accent_vietnamese(intern.name).upper()
            context['a5'] = intern.name_in_japan.replace(u'・', '  ')
            context['a7'] = intern_utils.date_time_in_jp(intern.day,intern.month,intern.year)
            context['a8'] = intern_utils.date_time_in_vn(intern.day,intern.month,intern.year)
            context['a11'] = str(intern_utils.get_age_jp(self.date_create_letter_promotion_short,intern.day, intern.month,intern.year))
            if intern.hktt:
                context['a15_1'] = intern_utils.no_accent_vietnamese(intern.hktt).upper()
                context['a15'] = intern_utils.convert_to_docx_string(intern.hktt).upper()

            context['a21'] = u'%s ~ %s'%(intern_utils.date_time_in_vn2(intern.last_education_from_month,intern.last_education_from_year),
                                        intern_utils.date_time_in_vn2( intern.last_education_to_month,
                                                                     intern.last_education_to_year))
            context['a21_1'] = u'%s ～ %s'%(intern_utils.date_time_in_jp(None,intern.last_education_from_month,intern.last_education_from_year),
                                          intern_utils.date_time_in_jp(None, intern.last_education_to_month,
                                                                       intern.last_education_to_year))

            context['a22'] = intern_utils.convert_to_docx_string(intern.last_school_education_jp)
            if intern.last_school_education:
                context['a23'] = intern_utils.convert_to_docx_string(intern.last_school_education)

            if intern.last_education_from_month2 and intern.last_education_from_year2 and \
                    intern.last_education_to_month2 and intern.last_education_to_year2 \
                     and intern.last_school_education_jp2 and intern.last_school_education2:
                context['b21'] = u'%s ~ %s' % (
                intern_utils.date_time_in_vn2( intern.last_education_from_month2, intern.last_education_from_year2),
                intern_utils.date_time_in_vn2( intern.last_education_to_month2,
                                             intern.last_education_to_year2))
                context['b21_1'] = u'%s ～ %s' % (
                intern_utils.date_time_in_jp(None, intern.last_education_from_month2, intern.last_education_from_year2),
                intern_utils.date_time_in_jp(None, intern.last_education_to_month2,
                                             intern.last_education_to_year2))
                context['b22'] = intern_utils.convert_to_docx_string(intern.last_school_education_jp2)
                if intern.last_school_education2:
                    context['b23'] = intern_utils.convert_to_docx_string(intern.last_school_education2)

            table_jobs = []
            for i in range(0,4):
                if i == 0:
                    info = {}
                    if intern.time_employee and intern.job_employee_jp:
                        info['a25'] = u'%s ~ %s' % (
                        intern_utils.date_time_in_vn2(intern.time_employee_from_month, intern.time_employee_from_year),
                        intern_utils.date_time_in_vn2(intern.time_employee_to_month,
                                                     intern.time_employee_to_year))
                        info['a25_1'] = intern.time_employee
                        info['a26'] = intern_utils.convert_to_docx_string(intern.job_employee_jp)
                        if intern.job_employee_vi:
                            info['a27'] = intern_utils.convert_to_docx_string(intern.job_employee_vi)
                        table_jobs.append(info)
                    else:
                        break
                if i == 1:
                    info = {}
                    if intern.time_employee2 and intern.job_employee2_jp:
                        info['a25'] = u'%s ~ %s' % (
                        intern_utils.date_time_in_vn2(intern.time_employee2_from_month, intern.time_employee2_from_year),
                        intern_utils.date_time_in_vn2(intern.time_employee2_to_month,
                                                     intern.time_employee2_to_year))
                        info['a25_1'] = intern.time_employee2
                        info['a26'] = intern_utils.convert_to_docx_string(intern.job_employee2_jp)
                        if intern.job_employee2_vi:
                            info['a27'] = intern_utils.convert_to_docx_string(intern.job_employee2_vi)
                        table_jobs.append(info)
                    else:
                        break
                if i == 2:
                    info = {}
                    if intern.time_employee3 and intern.job_employee3_jp :
                        info['a25'] = u'%s ~ %s' % (
                        intern_utils.date_time_in_vn2(intern.time_employee3_from_month, intern.time_employee3_from_year),
                        intern_utils.date_time_in_vn2(intern.time_employee3_to_month,
                                                     intern.time_employee3_to_year))
                        info['a25_1'] = intern.time_employee3
                        info['a26'] = intern_utils.convert_to_docx_string(intern.job_employee3_jp)
                        if intern.job_employee3_vi:
                            info['a27'] = intern_utils.convert_to_docx_string(intern.job_employee3_vi)
                        table_jobs.append(info)
                    else:
                        break
                if i == 3:
                    info = {}
                    if intern.time_employee4 and intern.job_employee4_jp:
                        info['a25'] = u'%s ~ %s' % (
                        intern_utils.date_time_in_vn2(intern.time_employee4_from_month, intern.time_employee4_from_year),
                        intern_utils.date_time_in_vn2(intern.time_employee4_to_month,
                                                     intern.time_employee4_to_year))
                        info['a25_1'] = intern.time_employee4
                        info['a26'] = intern_utils.convert_to_docx_string(intern.job_employee4_jp)
                        if intern.job_employee4_vi:
                            info['a27'] = intern_utils.convert_to_docx_string(intern.job_employee4_vi)
                        table_jobs.append(info)
                    else:
                        break

                if i == 4:
                    info = {}
                    if intern.time_employee5 and intern.job_employee5_jp:
                        info['a25'] = u'%s ~ %s' % (
                        intern_utils.date_time_in_vn2(intern.time_employee5_from_month, intern.time_employee5_from_year),
                        intern_utils.date_time_in_vn2(intern.time_employee5_to_month,
                                                     intern.time_employee5_to_year))
                        info['a25_1'] = intern.time_employee5
                        info['a26'] = intern_utils.convert_to_docx_string(intern.job_employee5_jp)
                        if intern.job_employee5_vi:
                            info['a27'] = intern_utils.convert_to_docx_string(intern.job_employee5_vi)
                        table_jobs.append(info)
                    else:
                        break


            info = {}

            info['a25'] = u'%s ~ Hiện nay' % (
                                intern_utils.date_time_in_vn2(intern.time_start_at_pc_from_month, intern.time_start_at_pc_from_year))
            info['a25_1'] = intern.time_start_at_pc

            # tmp = index/10
            # _logger.info("%d TMP "%tmp)
            info['a26'] = u"%s (%s)" %(intern_utils.convert_to_docx_string(intern.dispatchcom2.name),intern_utils.convert_to_docx_string(self.job_jp))
            if intern.dispatchcom2.name_vn:
                info['a27'] = u"%s (%s)" %(intern_utils.convert_to_docx_string(intern.dispatchcom2.name_vn),intern_utils.convert_to_docx_string(self.job_vi))
            table_jobs.append(info)
            if len(table_jobs) == 1:
                table_jobs.append({})

            context['tbl_jobs'] = table_jobs

            context['a38'] = u'%d年%dヶ月'%(intern.time_at_pc_year,intern.time_at_pc_month)
            context['a39'] = u'%d năm %d tháng'%(intern.time_at_pc_year,intern.time_at_pc_month)


            context['a40'] = intern_utils.date_time_in_jp(self.day_create_letter_promotion,
                                                          self.month_create_letter_promotion,
                                                          self.year_create_letter_promotion)

            context['a42'] = intern_utils.date_time_in_vn(self.day_create_letter_promotion,
                                                          self.month_create_letter_promotion,
                                                          self.year_create_letter_promotion)




            context['a84'] = intern_utils.convert_to_docx_string(self.job_jp)
            context['a85'] = intern_utils.convert_to_docx_string(self.job_vi)

            # _logger.info("AAAA %s" % str(context))

            tpl.render(context)


            tempFile = NamedTemporaryFile(delete=False)
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_doc_1_10(self, intern):
        docs = self.env['intern.document'].search([('name', '=', "Doc1-10")], limit=1)
        if docs:
            stream = BytesIO(docs[0].attachment.decode("base64"))
            tpl = DocxTemplate(stream)
            context = {}
            context['a3'] = intern_utils.no_accent_vietnamese(intern.name).upper()
            context['a59'] = intern_utils.convert_to_docx_string(self.guild.name_in_jp)
            context['a74'] = intern_utils.convert_to_docx_string(intern.enterprise.name_jp)
            context['a40'] = intern_utils.date_time_in_jp(self.day_create_letter_promotion,
                                                          self.month_create_letter_promotion,
                                                          self.year_create_letter_promotion)
            context['a102'] = intern_utils.convert_to_docx_string(self.dispatchcom1.name_en).upper()

            context['a102_1'] = intern_utils.convert_to_docx_string(self.dispatchcom1.name_en).upper()
            if '(' in context['a102_1'] and ')' in context['a102_1']:
                first = context['a102_1'].rfind('(')
                last = context['a102_1'].rfind(')')
                context['a102_1'] = context['a102_1'][:first] + context['a102_1'][last + 1:]


            context['a105'] = intern_utils.no_accent_vietnamese(self.dispatchcom1.director).upper()

            tpl.render(context)

            tempFile = NamedTemporaryFile(delete=False)
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_doc_1_20(self):
        if self.back_to_pc2:
            docs = self.env['intern.document'].search([('name', '=', "Doc1-20")], limit=1)
        else:
            docs = self.env['intern.document'].search([('name', '=', "Doc1-20-TH2")], limit=1)
        if docs:
            stream = BytesIO(docs[0].attachment.decode("base64"))
            tpl = DocxTemplate(stream)
            context = {}

            context['a40'] = intern_utils.date_time_in_jp(self.day_create_letter_promotion,
                                                          self.month_create_letter_promotion,
                                                          self.year_create_letter_promotion)
            context['a42'] = intern_utils.date_time_in_vn(self.day_create_letter_promotion,
                                                          self.month_create_letter_promotion,
                                                          self.year_create_letter_promotion)

            context['a84'] = intern_utils.convert_to_docx_string(self.job_jp)
            context['a85'] = intern_utils.convert_to_docx_string(self.job_vi).lower()
            tpl.render(context)

            tempFile = NamedTemporaryFile(delete=False)
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_doc_1_21(self, intern):
        if self.year_expire == 3:
            docs = self.env['intern.document'].search([('name', '=', "Doc1-21")], limit=1)
        else:
            docs = self.env['intern.document'].search([('name', '=', "Doc1-21-1")], limit=1)
        if docs:
            stream = BytesIO(docs[0].attachment.decode("base64"))
            tpl = DocxTemplate(stream)
            context = {}

            context['a59'] = intern_utils.convert_to_docx_string(self.guild.name_in_jp)
            context['a60'] = intern_utils.convert_to_docx_string(self.guild.name_in_en)
            context['a74'] = intern_utils.convert_to_docx_string(intern.enterprise.name_jp)
            context['a75'] = intern_utils.convert_to_docx_string(intern.enterprise.name_romaji)
            context['a3'] = intern_utils.no_accent_vietnamese(intern.name).upper()
            context['a5'] = intern.name_in_japan.replace(u'・', '  ')
            context['a47'] = intern_utils.date_time_in_jp(self.day_pay_finance1,self.month_pay_finance1,self.year_pay_finance1)
            context['a48'] = self.date_pay_finance1
            context['a49'] = intern_utils.date_time_in_jp(self.day_pay_finance2, self.month_pay_finance2,
                                                          self.year_pay_finance2)
            context['a50'] = self.date_pay_finance2

            context['a40'] = intern_utils.date_time_in_jp(self.day_create_letter_promotion,
                                                          self.month_create_letter_promotion,
                                                          self.year_create_letter_promotion)
            context['a42'] = intern_utils.date_time_in_vn(self.day_create_letter_promotion,
                                                          self.month_create_letter_promotion,
                                                          self.year_create_letter_promotion)

            context['a102'] = intern_utils.convert_to_docx_string(self.dispatchcom1.name_en).upper()

            context['a102_1'] = intern_utils.convert_to_docx_string(self.dispatchcom1.name_en).upper()
            if '(' in context['a102_1']:
                first = context['a102_1'].rfind('(')
                last = context['a102_1'].rfind(')')
                context['a102_1'] = context['a102_1'][:first] + context['a102_1'][last + 1:]

            context['a104'] = intern_utils.convert_to_docx_string(self.dispatchcom1.name).upper()
            context['a105'] = intern_utils.no_accent_vietnamese(self.dispatchcom1.director).upper()
            context['a105_1'] = self.dispatchcom1.director.upper()
            context['a106'] = intern_utils.convert_to_docx_string(self.dispatchcom1.position_director)
            context['a123'] = intern_utils.convert_to_docx_string(self.dispatchcom1.position_director_vi)

            context['a72'] = "{:,}".format(self.guild.fee_training_nd_to_pc)
            context['a72x2'] = "{:,}".format(200*self.guild.fee_training_nd_to_pc)

            tpl.render(context)

            tempFile = NamedTemporaryFile(delete=False)
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_doc_1_28(self, intern, index):
        if self.back_to_pc2:
            docs = self.env['intern.document'].search([('name', '=', "Doc1-28")], limit=1)
        else:
            docs = self.env['intern.document'].search([('name', '=', "Doc1-28-TH2")], limit=1)
        if docs:
            stream = BytesIO(docs[0].attachment.decode("base64"))
            tpl = DocxTemplate(stream)
            context = {}

            context['a3'] = intern_utils.no_accent_vietnamese(intern.name).upper()
            context['a5'] = intern.name_in_japan.replace(u'・', '  ')
            context['a84'] = intern_utils.convert_to_docx_string(self.job_jp)
            context['a82'] = intern_utils.convert_to_docx_string(self.name_working_department)
            # tmp = index/10
            context['a92'] = intern_utils.convert_to_docx_string(intern.dispatchcom2.name)
            context['a94'] = intern.dispatchcom2.position_person_sign
            context['a95'] = intern.dispatchcom2.director.upper()
            context['a40'] = intern_utils.date_time_in_jp(self.day_create_letter_promotion,self.month_create_letter_promotion,
                                                          self.year_create_letter_promotion)

            if self.back_to_pc2:
                context['a92'] = intern_utils.convert_to_docx_string(intern.dispatchcom2.name)

            tpl.render(context)

            tempFile = NamedTemporaryFile(delete=False)
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile




    def create_hdtn(self, intern):
        docs = self.env['intern.document'].search([('name', '=', "DocHDTN")], limit=1)
        if docs:
            stream = BytesIO(docs[0].attachment.decode("base64"))
            tpl = DocxTemplate(stream)
            context = {}

            context['a3'] = intern_utils.no_accent_vietnamese(intern.name).upper()
            context['a7'] = intern_utils.date_time_in_jp(intern.day,intern.month,intern.year)

            context['a10'] = intern_utils.date_time_in_jp(intern.day_identity,intern.month_identity,intern.year_identity)
            context['a15'] = intern_utils.no_accent_vietnamese(intern.hktt).upper()
            if intern.identity:
                context['a16'] = u'%s省公安'%(intern_utils.no_accent_vietnamese(intern.place_cmnd.name).upper())
                context['a9'] = intern.identity
            else:
                context['a9'] = intern.identity_2
                context['a16'] = u'警察局局長'

            context['a17'] = intern_utils.no_accent_vietnamese(intern.contact_person).upper()
            context['a18'] = intern_utils.convert_to_docx_string(intern_utils.no_accent_vietnamese(intern.contact_address)).upper()
            context['a19'] = intern.contact_relative.relation_jp
            context['a20'] = intern.contact_phone
            context['a59'] = self.guild.name_in_jp
            context['a70'] = intern_utils.date_time_in_jp(self.guild.day_sign,self.guild.month_sign,self.guild.year_sign)

            if self.guild.note_subsize_jp:
                if not self.guild.subsidize_start_month or self.guild.subsidize_start_month == 0:
                    context['a73'] = u'%s'%self.guild.note_subsize_jp
                else:
                    context['a73'] = u'%s 円 (%s)'%(str("{:,}".format(self.guild.subsidize_start_month)),self.guild.note_subsize_jp)
            else:
                context['a73'] = u'%s 円'%str("{:,}".format(self.guild.subsidize_start_month))

            context['a74'] = intern_utils.convert_to_docx_string(intern.enterprise.name_jp)
            context['a76'] = intern_utils.convert_to_docx_string(intern.enterprise.address_jp)
            context['a77'] = intern.enterprise.phone_number
            context['a84'] = intern_utils.convert_to_docx_string(self.job_jp)
            context['a86'] = self.year_expire

            context['a122'] = intern_utils.date_time_in_jp_missing(self.day_sign_proletter,self.month_sign_proletter,self.year_sign_proletter)
            context['a122_3'] = self.year_sign_proletter

            context['a102'] = intern_utils.convert_to_docx_string(self.dispatchcom1.name_en).upper()
            context['a105'] = intern_utils.no_accent_vietnamese(self.dispatchcom1.director).upper()
            context['a106'] = intern_utils.convert_to_docx_string(self.dispatchcom1.position_director)
            context['a112'] = intern_utils.convert_to_docx_string(self.dispatchcom1.address_en)
            context['a110'] = self.dispatchcom1.phone_number
            if self.dispatchcom1.fax_number:
                context['a111'] = self.dispatchcom1.fax_number

            tpl.render(context)

            tempFile = NamedTemporaryFile(delete=False)
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_hdtv(self, intern):
        docs = self.env['intern.document'].search([('name', '=', "DocHDTV")], limit=1)
        if docs:
            stream = BytesIO(docs[0].attachment.decode("base64"))
            tpl = DocxTemplate(stream)
            context = {}

            context['a1'] = intern.name.upper()
            context['a6'] = intern_utils.date_time_in_en(intern.day, intern.month, intern.year)

            context['a10_1'] = intern_utils.date_time_in_en(intern.day_identity, intern.month_identity,
                                                          intern.year_identity)
            context['a15_1'] = intern.hktt
            if intern.identity:
                context['a16_1'] = u'Công An %s' %(intern.place_cmnd.name)
                context['a9'] = intern.identity
            else:
                context['a9'] = intern.identity_2
                context['a16_1'] = u'Cục trưởng Cục cảnh sát'

            context['a17_1'] = intern.contact_person
            context['a18_1'] = intern.contact_address
            context['a19_1'] = intern.contact_relative.relation
            context['a20'] = intern.contact_phone
            context['a60_2'] = intern_utils.convert_to_docx_string(self.guild.name_in_en.upper()).replace('KYODO KUMIAI','').replace('KYOUDOU KUMIAI','')
            context['a71'] = intern_utils.date_time_in_en(self.guild.day_sign, self.guild.month_sign,
                                                          self.guild.year_sign)

            if self.guild.note_subsize_vi:
                if not self.guild.subsidize_start_month or self.guild.subsidize_start_month == 0:
                    context['a73'] = u'%s'%self.guild.note_subsize_vi
                else:
                    context['a73'] = u'%s Yên (%s)'%(intern_utils.format_number_in_vn(str(self.guild.subsidize_start_month)),self.guild.note_subsize_vi)
            else:
                context['a73'] = u'%s Yên'%intern_utils.format_number_in_vn(str(self.guild.subsidize_start_month))

            context['a75'] = intern_utils.convert_to_docx_string(intern.enterprise.name_romaji).upper()
            context['a79'] = intern_utils.convert_to_docx_string(intern.enterprise.address_romoji)
            context['a77'] = intern.enterprise.phone_number
            context['a85'] = intern_utils.convert_to_docx_string(self.job_vi)
            context['a86'] = self.year_expire

            context['a122'] = intern_utils.date_time_in_en_missing(self.day_sign_proletter, self.month_sign_proletter,
                                                           self.year_sign_proletter)
            context['a122_3'] = self.year_sign_proletter
            context['a104'] = intern_utils.convert_to_docx_string(self.dispatchcom1.name)
            context['a105'] = self.dispatchcom1.director
            context['a123'] = intern_utils.convert_to_docx_string(self.dispatchcom1.position_director_vi)
            context['a124'] = intern_utils.convert_to_docx_string(self.dispatchcom1.address_vi)
            context['a110'] = self.dispatchcom1.phone_number
            if self.dispatchcom1.fax_number:
                context['a111'] = self.dispatchcom1.fax_number

            tpl.render(context)

            tempFile = NamedTemporaryFile(delete=False)
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile


    def create_doc_1_29(self,enterprise_id):
        docs = self.env['intern.document'].search([('name', '=', "Doc1-29")], limit=1)
        if docs:
            stream = BytesIO(docs[0].attachment.decode("base64"))
            tpl = DocxTemplate(stream)
            context = {}

            table_interns = []
            if not self.hoso_created:
                interns_pass = sorted(self.interns_pass_doc, key=lambda x: x.sequence_pass)
            else:
                interns_pass = sorted(self.interns_pass_doc_hs, key=lambda x: x.sequence_pass)
            counter = 0
            for i, intern in enumerate(interns_pass):
                if intern.enterprise.id != enterprise_id:
                    continue
                counter+=1
                info = {}
                info['stt'] = str(counter)
                info['htk'] = intern_utils.no_accent_vietnamese(intern.name).upper()
                info['xc'] = intern_utils.date_time_in_jp(self.day_departure_doc,self.month_departure_doc,self.year_departure_doc)

                table_interns.append(info)


            context['a43'] = intern_utils.date_time_in_jp(self.day_create_plan_training,self.month_create_plan_training,self.year_create_plan_training)
            context['a44'] = intern_utils.date_time_in_jp(self.day_start_training,self.month_start_training,self.year_start_training)
            context['a45'] = intern_utils.date_time_in_jp(self.day_end_training,self.month_end_training,self.year_end_training)
            context['a53'] = intern_utils.convert_to_docx_string(self.training_center.name_jp)
            # context['a54'] = self.training_center.address_jp
            context['a54'] = intern_utils.convert_to_docx_string(self.training_center.address_en)
            context['a59'] = intern_utils.convert_to_docx_string(self.guild.name_in_jp)
            context['a61'] = intern_utils.convert_to_docx_string(self.guild.address_in_jp)
            context['a66'] = intern_utils.convert_to_docx_string(self.guild.position_of_responsive_jp)
            context['a68'] = self.guild.name_of_responsive_jp

            context['tbl_intern'] = table_interns

            tpl.render(context)

            tempFile = NamedTemporaryFile(delete=False)
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_master(self, enterprise_id):
        docs = self.env['intern.document'].search([('name', '=', "DocMaster")], limit=1)

        list_interns = self.interns_pass_doc
        if self.hoso_created:
            list_interns = self.interns_pass_doc_hs

        interns_pass = sorted(list_interns, key=lambda x: x.sequence_pass)
        if docs:
            stream = BytesIO(docs[0].attachment.decode("base64"))
            tpl = DocxTemplate(stream)
            context = {}
            intern = None
            for ite in interns_pass:
                if ite.enterprise.id == enterprise_id:
                    intern = ite
                    break
            counter_intern_base_enterprise = 0
            for intern in interns_pass:
                if ite.enterprise.id == enterprise_id:
                    counter_intern_base_enterprise+=1
            context['a1'] = intern.name.upper()
            context['a2'] = intern.name.upper()
            context['a3'] = intern_utils.no_accent_vietnamese(intern.name).upper()
            context['a4'] = str(counter_intern_base_enterprise-1)
            context['a5'] = intern.name_in_japan.replace(u'・', ' ')
            context['a6'] = intern_utils.date_time_in_en(intern.day, intern.month, intern.year)
            context['a7'] = intern_utils.date_time_in_jp(intern.day,intern.month,intern.year)
            context['a8'] = intern_utils.date_time_in_vn(intern.day, intern.month, intern.year)
            if intern.identity:
                context['a9'] = intern.identity
                context['a16'] = u'%s省公安' % (intern_utils.no_accent_vietnamese(intern.place_cmnd.name))
            else:
                context['a9'] = intern.identity_2
                context['a16'] = u'警察局局長'
            context['a10'] = intern_utils.date_time_in_jp(intern.day_identity, intern.month_identity,
                                                          intern.year_identity)
            context['a11'] = str(intern_utils.get_age_jp(self.date_create_letter_promotion_short,intern.day,intern.month,intern.year))
            if intern.gender == 'nam':
                context['a12'] = u'男'
                context['a13'] = u'MALE'
            else:
                context['a12'] = u'女'
                context['a13'] = u'FEMALE'
            if intern.marital_status.id is not 2:
                context['a14'] = u'無'
            else:
                context['a14'] = u'有'

            if intern.hktt:
                context['a15'] = intern_utils.no_accent_vietnamese(intern.hktt).upper()

            context['a17'] = intern_utils.no_accent_vietnamese(intern.contact_person).upper()
            context['a18'] = intern_utils.no_accent_vietnamese(intern.contact_address).upper()
            context['a19'] = intern.contact_relative.relation_jp
            context['a20'] = intern.contact_phone
            context['a21'] = intern.last_time_education
            context['a22'] = intern_utils.convert_to_docx_string(intern.last_school_education_jp)
            context['a23'] = intern_utils.convert_to_docx_string(intern.last_school_education)
            context['a24'] = intern.last_time_education
            context['a25'] = intern.time_employee
            context['a26'] = intern_utils.convert_to_docx_string(intern.job_employee_jp)
            context['a27'] = intern_utils.convert_to_docx_string(intern.job_employee_vi)

            if intern.time_employee2 and intern.job_employee2_jp and intern.job_employee2_vi:
                context['a28'] = intern.time_employee2
                context['a29'] = intern_utils.convert_to_docx_string(intern.job_employee2_jp)
                context['a30'] = intern_utils.convert_to_docx_string(intern.job_employee2_vi)

            if intern.time_employee3 and intern.job_employee3_jp and intern.job_employee3_vi:
                context['a31'] = intern.time_employee3
                context['a32'] = intern_utils.convert_to_docx_string(intern.job_employee3_jp)
                context['a33'] = intern_utils.convert_to_docx_string(intern.job_employee3_vi)

            context['a34'] = intern_utils.date_time_in_jp(intern.dispatchcom2.day_create, intern.dispatchcom2.month_create,
                                                          intern.dispatchcom2.year_create)
            context['a35'] = intern.time_start_at_pc
            context['a36'] = "%s (%s)" % (intern_utils.convert_to_docx_string(intern.dispatchcom2.name), intern_utils.convert_to_docx_string(self.job_jp))
            context['a37'] = "%s (%s)" % (intern_utils.convert_to_docx_string(intern.dispatchcom2.name_vn), intern_utils.convert_to_docx_string(self.job_vi))
            context['a38'] = u'%d年%dヶ月' % (intern.time_at_pc_year, intern.time_at_pc_month)
            context['a39'] = u'%d năm %d tháng' % (intern.time_at_pc_year, intern.time_at_pc_month)

            context['a40'] = intern_utils.date_time_in_jp(self.day_create_letter_promotion,
                                                          self.month_create_letter_promotion,
                                                          self.year_create_letter_promotion)
            context['a41'] = intern_utils.date_time_in_en(self.day_create_letter_promotion,
                                                          self.month_create_letter_promotion,
                                                          self.year_create_letter_promotion)

            context['a42'] = intern_utils.date_time_in_vn(self.day_create_letter_promotion,
                                                          self.month_create_letter_promotion,
                                                          self.year_create_letter_promotion)

            context['a43'] = intern_utils.date_time_in_jp(self.day_create_plan_training,
                                                          self.month_create_plan_training,
                                                          self.year_create_plan_training)
            context['a44'] = intern_utils.date_time_in_jp(self.day_start_training, self.month_start_training,
                                                          self.year_start_training)
            context['a45'] = intern_utils.date_time_in_jp(self.day_end_training, self.month_end_training,
                                                          self.year_end_training)

            context['a46'] = intern_utils.date_time_in_jp(self.day_create_plan_training_report_customer,self.month_create_plan_training_report_customer,self.year_create_plan_training_report_customer)
            context['a47'] = intern_utils.date_time_in_jp(self.day_pay_finance1, self.month_pay_finance1,
                                                          self.year_pay_finance1)
            context['a48'] = self.date_pay_finance1
            context['a49'] = intern_utils.date_time_in_jp(self.day_pay_finance2, self.month_pay_finance2,
                                                          self.year_pay_finance2)
            context['a50'] = self.date_pay_finance2
            context['a51'] = self.length_training
            context['a52'] = str(self.hours_training)
            context['a53'] = intern_utils.convert_to_docx_string(self.training_center.name_jp)
            # context['a54'] = self.training_center.address_jp
            context['a55'] = intern_utils.date_time_in_jp(self.training_center.day_create,self.training_center.month_create,self.training_center.year_create)
            context['a56'] = self.training_center.phone_number
            context['a57'] = intern_utils.no_accent_vietnamese(self.training_center.responsive_person)
            context['a58'] = intern_utils.convert_to_docx_string(self.guild.name_acronym)
            context['a59'] = intern_utils.convert_to_docx_string(self.guild.name_in_jp)
            context['a60'] = intern_utils.convert_to_docx_string(self.guild.name_in_en).upper()
            context['a61'] = intern_utils.convert_to_docx_string(self.guild.address_in_jp)
            context['a62'] = intern_utils.convert_to_docx_string(self.guild.address_in_romaji)
            context['a63'] = self.guild.post_code
            context['a64'] = self.guild.phone_number
            if self.guild.fax_number:
                context['a65'] = self.guild.fax_number
            context['a66'] = intern_utils.convert_to_docx_string(self.guild.position_of_responsive_jp)
            context['a67'] = intern_utils.convert_to_docx_string(self.guild.position_of_responsive_vi)
            context['a68'] = self.guild.name_of_responsive_jp
            context['a69'] = self.guild.name_of_responsive_romaji
            context['a70'] = intern_utils.date_time_in_jp(self.guild.day_sign, self.guild.month_sign,
                                                          self.guild.year_sign)
            context['a71'] = intern_utils.date_time_in_en(self.guild.day_sign, self.guild.month_sign,
                                                          self.guild.year_sign)
            context['a72'] = self.guild.fee_training_nd_to_pc
            context['a73'] = str(self.guild.subsidize_start_month)
            context['a74'] = intern_utils.convert_to_docx_string(intern.enterprise.name_jp)
            context['a75'] = intern_utils.convert_to_docx_string(intern.enterprise.name_romaji)
            context['a76'] = intern_utils.convert_to_docx_string(intern.enterprise.address_jp)
            context['a77'] = intern.enterprise.phone_number
            if intern.enterprise.fax_number:
                context['a78'] = intern.enterprise.fax_number
            context['a79'] = intern_utils.convert_to_docx_string(intern.enterprise.address_romoji)
            context['a80'] = intern.enterprise.name_of_responsive_jp
            context['a81'] = intern.enterprise.name_of_responsive_en
            context['a82'] = intern_utils.convert_to_docx_string(self.name_working_department)
            context['a83'] = intern_utils.convert_to_docx_string(self.job_en)
            context['a84'] = intern_utils.convert_to_docx_string(self.job_jp)
            context['a85'] = intern_utils.convert_to_docx_string(self.job_vi)
            context['a86'] = self.year_expire
            context['a87'] = intern_utils.date_time_in_jp(self.day_departure_doc,self.month_departure_doc,self.year_departure_doc)
            context['a88'] = self.year_departure_doc
            context['a89'] = self.month_departure_doc
            context['a90'] = self.person_sign_proletter
            context['a91'] = intern_utils.convert_to_docx_string(self.position_person_sign)
            context['a92'] = intern_utils.convert_to_docx_string(intern.dispatchcom2.name)
            if intern.dispatchcom2.address:
                context['a93'] = intern_utils.convert_to_docx_string(intern.dispatchcom2.address)
            context['a95'] = intern_utils.no_accent_vietnamese(intern.dispatchcom2.director).upper()
            context['a96'] = intern_utils.convert_to_docx_string(intern.dispatchcom2.position_person_sign)
            context['a97'] = intern.dispatchcom2.phone_number
            if intern.dispatchcom2.fax_number:
                context['a98'] = intern.dispatchcom2.fax_number
            context['a99'] = intern_utils.date_time_in_jp(intern.dispatchcom2.day_create,intern.dispatchcom2.month_create,intern.dispatchcom2.year_create)
            context['a100'] = self.developing_employee

            #Phap nhan
            context['a101'] = intern_utils.convert_to_docx_string(self.dispatchcom1.name_jp)
            context['a102'] = intern_utils.convert_to_docx_string(self.dispatchcom1.name_en)
            context['a104'] = intern_utils.convert_to_docx_string(self.dispatchcom1.name)
            context['a105'] = intern_utils.no_accent_vietnamese(self.dispatchcom1.director).upper()
            context['a106'] = self.dispatchcom1.position_director
            # context['a107'] = self.dispatchcom1.address_jp
            context['a108'] = intern_utils.convert_to_docx_string(self.dispatchcom1.address_en)
            context['a110'] = self.dispatchcom1.phone_number
            if self.dispatchcom1.fax_number:
                context['a111'] = self.dispatchcom1.fax_number

            context['a112'] = intern_utils.convert_to_docx_string(self.dispatchcom1.address_en)
            context['a113'] = intern_utils.date_time_in_jp(self.dispatchcom1.day_create,self.dispatchcom1.month_create,self.dispatchcom1.year_create)

            counter = 113
            #TTS di cung
            if counter_intern_base_enterprise >1:
                table_interns = []
                iterate_intern = 0
                for i, itern in enumerate(interns_pass):
                    if itern.enterprise.id!=enterprise_id:
                        continue

                    if iterate_intern>0:
                        info = {}
                        info['a114'] = intern_utils.no_accent_vietnamese(itern.name).upper()
                        info['a115'] = intern_utils.no_accent_vietnamese(itern.name).upper()
                        info['a116'] = itern.name_in_japan.replace(u'・', ' ')
                        info['a117'] = intern_utils.date_time_in_en(itern.day, itern.month, itern.year)
                        info['a118'] = intern_utils.date_time_in_jp(itern.day, itern.month, itern.year)
                        info['a119'] = intern_utils.date_time_in_en(itern.day, itern.month, itern.year)
                        info['a120'] = intern_utils.date_time_in_vn(itern.day, itern.month, itern.year)
                        info['a121'] = str(intern_utils.get_age_jp(self.date_create_letter_promotion_short,itern.day, itern.month,itern.year))


                        if itern.gender == 'nam':
                            info['a122'] = u'男'
                            info['a123'] = u'MALE'
                        else:
                            info['a122'] = u'女'
                            info['a123'] = u'FEMALE'
                        if itern.marital_status.id is not 2:
                            info['a124'] = u'無'
                        else:
                            info['a124'] = u'有'

                        if itern.identity:
                            info['a125'] = itern.identity
                            info['a127'] = u'%s省公安' % (intern_utils.no_accent_vietnamese(itern.place_cmnd.name))
                        else:
                            info['a15'] = itern.identity_2
                            info['a127'] = u'警察局局長'

                        info['a126'] = intern_utils.date_time_in_jp(itern.day_identity, itern.month_identity,
                                                                    itern.year_identity)

                        if itern.hktt:
                            info['a128'] = intern_utils.no_accent_vietnamese(itern.hktt).upper()
                        info['a129'] = itern.last_time_education
                        info['a130'] = intern_utils.convert_to_docx_string(itern.last_school_education_jp)
                        info['a131'] = intern_utils.convert_to_docx_string(itern.last_school_education)
                        info['a132'] = itern.last_time_education

                        info['a133'] = itern.time_employee
                        info['a134'] = intern_utils.convert_to_docx_string(itern.job_employee_jp)
                        info['a135'] = intern_utils.convert_to_docx_string(itern.job_employee_vi)

                        if itern.time_employee2 and itern.job_employee2_jp and itern.job_employee2_vi:
                            info['a136'] = itern.time_employee2
                            info['a137'] = intern_utils.convert_to_docx_string(itern.job_employee2_jp)
                            info['a138'] = intern_utils.convert_to_docx_string(itern.job_employee2_vi)

                        if itern.time_employee3 and itern.job_employee3_jp and itern.job_employee3_vi:
                            info['a139'] = itern.time_employee3
                            info['a140'] = intern_utils.convert_to_docx_string(itern.job_employee3_jp)
                            info['a141'] = intern_utils.convert_to_docx_string(itern.job_employee3_vi)

                        # tmp = iterate_intern/10
                        info['a142'] = intern_utils.date_time_in_jp(itern.dispatchcom2.day_create,itern.dispatchcom2.month_create,itern.dispatchcom2.year_create)
                        info['a143'] = itern.time_start_at_pc
                        info['a144'] = u'%s (%s)' % (intern_utils.convert_to_docx_string(itern.dispatchcom2.name), intern_utils.convert_to_docx_string(self.job_jp))
                        info['a145'] = u'%s (%s)' % (intern_utils.convert_to_docx_string(itern.dispatchcom2.name_vn), intern_utils.convert_to_docx_string(self.job_vi))
                        info['a146'] = u'%d年%dヶ月' % (itern.time_at_pc_year, itern.time_at_pc_month)
                        info['a147'] = u'%d năm %d tháng' % (itern.time_at_pc_year, itern.time_at_pc_month)
                        # _logger.info(u"itern.contact_person %s %s" %(itern.contact_person,itern.name))
                        info['a148'] = intern_utils.no_accent_vietnamese(itern.contact_person).upper()
                        info['a149'] = intern_utils.no_accent_vietnamese(itern.contact_address).upper()
                        info['a150'] = itern.contact_relative.relation_jp
                        info['a151'] = itern.contact_phone
                        #stt
                        info['stt'] = str(iterate_intern+1)
                        for x in range(1,39):
                            info['s%d'%x] =str(counter+x)
                        counter = counter+38
                        table_interns.append(info)
                    iterate_intern+=1

                context['tbl_intern'] = table_interns

            tpl.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_list_of_sent_en(self,enterprise_id):
        enterprise = self.env['intern.enterprise'].browse(int(enterprise_id))
        docs = self.env['intern.document'].search([('name', '=', "list_of_sent")], limit=1)
        if docs:
            stream = BytesIO(docs[0].attachment.decode("base64"))
            tpl = DocxTemplate(stream)
            context = {}
            context['a60'] = intern_utils.convert_to_docx_string(self.guild.name_in_en).upper()
            context['a69'] = intern_utils.convert_to_docx_string(self.guild.name_of_responsive_romaji)
            context['a61'] = intern_utils.convert_to_docx_string(self.guild.address_in_jp)
            context['a62'] = intern_utils.convert_to_docx_string(self.guild.address_in_romaji)
            if self.guild.license_number:
                context['a120'] = self.guild.license_number
            if self.guild.fax_number:
                context['a64_65'] = "%s/%s"%(self.guild.phone_number,self.guild.fax_number)
            else:
                context['a64_65'] = self.guild.phone_number

            context['a74'] = intern_utils.convert_to_docx_string(enterprise.name_jp)
            context['a75'] = intern_utils.convert_to_docx_string(enterprise.name_romaji)
            context['a81'] = enterprise.name_of_responsive_en
            context['a79'] = intern_utils.convert_to_docx_string(enterprise.address_romoji)
            context['a77'] = enterprise.phone_number
            context['a86'] = self.year_expire

            table_interns = []

            list_interns = self.interns_pass_doc
            if self.hoso_created:
                list_interns = self.interns_pass_doc_hs

            interns_pass = sorted(list_interns, key=lambda x: x.sequence_pass)
            counter=0
            for i, intern in enumerate(interns_pass):
                if intern.enterprise.id != int(enterprise_id):
                    continue
                counter+=1
                info = {}
                info['stt'] = str(counter)
                info['a3'] = intern_utils.no_accent_vietnamese(intern.name).upper()
                info['a6'] = intern_utils.date_time_in_en(intern.day, intern.month, intern.year)
                if intern.gender == 'nam':
                    info['a13'] = 'MALE'
                else:
                    info['a13'] = 'FEMALE'
                info['a83'] = intern_utils.convert_to_docx_string(self.job_en).upper()
                info['a88'] = self.year_departure_doc
                info['a89'] = self.month_departure_doc
                table_interns.append(info)
            context['tbl_intern'] = table_interns
            context['a41'] = intern_utils.date_time_in_en(self.day_create_letter_promotion,
                                                          self.month_create_letter_promotion,
                                                          self.year_create_letter_promotion)

            context['a91'] = self.position_person_sign
            context['a90'] = intern_utils.no_accent_vietnamese(self.person_sign_proletter).upper()

            context['a102'] = intern_utils.convert_to_docx_string(self.dispatchcom1.name_en)
            context['a121'] = self.dispatchcom1.license_number
            context['a105'] = intern_utils.no_accent_vietnamese(self.dispatchcom1.director).upper()
            context['a108'] = intern_utils.convert_to_docx_string(self.dispatchcom1.address_en)

            phone_tax = self.dispatchcom1.phone_number
            if self.dispatchcom1.fax_number and self.dispatchcom1.fax_number != self.dispatchcom1.phone_number:
                phone_tax = phone_tax + "/" + self.dispatchcom1.fax_number
            context['a110_111'] = phone_tax

            tpl.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_list_of_sent_jp(self,enterprise_id):
        enterprise = self.env['intern.enterprise'].browse(int(enterprise_id))
        docs = self.env['intern.document'].search([('name', '=', "list_of_sent_jp")], limit=1)
        if docs:
            stream = BytesIO(docs[0].attachment.decode("base64"))
            tpl = DocxTemplate(stream)
            context = {}
            context['a59'] = intern_utils.convert_to_docx_string(self.guild.name_in_jp)
            context['a60'] = intern_utils.convert_to_docx_string(self.guild.name_in_en).upper()
            context['a68'] = intern_utils.convert_to_docx_string(self.guild.name_of_responsive_jp)
            context['a62'] = intern_utils.convert_to_docx_string(self.guild.address_in_romaji)
            context['a61'] = intern_utils.convert_to_docx_string(self.guild.address_in_jp)
            if self.guild.license_number:
                context['a120'] = self.guild.license_number
            if self.guild.fax_number:
                context['a64_65'] = "%s/%s"%(self.guild.phone_number,self.guild.fax_number)
            else:
                context['a64_65'] = self.guild.phone_number

            context['a74'] = intern_utils.convert_to_docx_string(enterprise.name_jp)
            context['a75'] = intern_utils.convert_to_docx_string(enterprise.name_romaji)
            context['a80'] = intern_utils.convert_to_docx_string(enterprise.name_of_responsive_jp)
            context['a76'] = intern_utils.convert_to_docx_string(enterprise.address_jp)
            context['a77'] = enterprise.phone_number
            context['a86'] = self.year_expire

            table_interns = []
            list_interns = self.interns_pass_doc
            if self.hoso_created:
                list_interns = self.interns_pass_doc_hs

            interns_pass = sorted(list_interns, key=lambda x: x.sequence_pass)
            counter = 0
            for i, intern in enumerate(interns_pass):
                if intern.enterprise.id != int(enterprise_id):
                    continue
                counter+=1
                info = {}
                info['stt'] = str(counter)
                info['a3'] = intern_utils.no_accent_vietnamese(intern.name).upper()
                info['a7'] = intern_utils.date_time_in_jp(intern.day,intern.month,intern.year)
                if intern.gender == 'nam':
                    info['a12'] = u'男'
                else:
                    info['a12'] = u'女'
                info['a84'] = intern_utils.convert_to_docx_string(self.job_jp)
                info['a88'] = self.year_departure_doc
                info['a89'] = self.month_departure_doc
                table_interns.append(info)


            context['tbl_intern'] = table_interns
            context['a40'] = intern_utils.date_time_in_jp(self.day_create_letter_promotion,
                                                          self.month_create_letter_promotion,
                                                          self.year_create_letter_promotion)

            context['a91_1'] = self.position_person_sign_jp
            context['a90_1'] = intern_utils.no_accent_vietnamese(self.person_sign_proletter).upper()

            context['a102'] = intern_utils.convert_to_docx_string(self.dispatchcom1.name_en)
            context['a121'] = self.dispatchcom1.license_number
            context['a105'] = intern_utils.no_accent_vietnamese(self.dispatchcom1.director).upper()
            context['a108'] = intern_utils.convert_to_docx_string(self.dispatchcom1.address_en)

            phone_tax = self.dispatchcom1.phone_number
            if self.dispatchcom1.fax_number and self.dispatchcom1.fax_number != self.dispatchcom1.phone_number:
                phone_tax = phone_tax+"/"+self.dispatchcom1.fax_number
            context['a110_111'] = phone_tax

            tpl.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_1_13_1(self):
        doc1_13_1 = self.env['intern.document'].search([('name', '=', "Doc1-13-1")], limit=1)
        if doc1_13_1:
            stream = BytesIO(doc1_13_1[0].attachment.decode("base64"))
            tpl = DocxTemplate(stream)
            tempFile = NamedTemporaryFile(delete=False)
            context = {}
            context['a53'] = intern_utils.convert_to_docx_string(self.training_center.name_jp)
            context['a57'] = intern_utils.no_accent_vietnamese(self.training_center.responsive_person)
            context['a55'] = intern_utils.date_time_in_jp(self.training_center.day_create,
                                                          self.training_center.month_create,
                                                          self.training_center.year_create)

            context['a56'] = intern_utils.convert_to_vn_phone(self.training_center.phone_number)
            context['a125'] = intern_utils.convert_to_docx_string(self.training_center.address_en)

            mission = self.training_center.mission
            if '\n' in self.training_center.mission:
                pre = '<w:p><w:r><w:t>'
                post = '</w:t></w:r></w:p>'
                lineBreak = '<w:br/>'
                test = self.training_center.mission.replace('\n',lineBreak)
                mission = pre + test + post
            context['a126'] = intern_utils.convert_to_docx_string(mission)
            context['a127'] = self.training_center.number_of_employee

            context['a40'] = intern_utils.date_time_in_jp(self.day_create_letter_promotion,
                                                          self.month_create_letter_promotion,
                                                          self.year_create_letter_promotion)

            tpl.render(context)
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_1_13_2(self):
        doc1_13_1 = self.env['intern.document'].search([('name', '=', "Doc1-13-2")], limit=1)
        if doc1_13_1:
            stream = BytesIO(doc1_13_1[0].attachment.decode("base64"))
            tpl = DocxTemplate(stream)
            tempFile = NamedTemporaryFile(delete=False)
            context = {}
            context['a102'] = intern_utils.convert_to_docx_string(self.dispatchcom1.name_en).upper()
            context['a102_1'] = intern_utils.convert_to_docx_string(self.dispatchcom1.name_en).upper()
            if '(' in context['a102_1']:
                first = context['a102_1'].rfind('(')
                last = context['a102_1'].rfind(')')
                context['a102_1'] = context['a102_1'][:first] + context['a102_1'][last+1:]
            context['a105'] = intern_utils.no_accent_vietnamese(self.dispatchcom1.director)

            context['a113'] = intern_utils.date_time_in_jp(self.dispatchcom1.day_create,
                                                          self.dispatchcom1.month_create,
                                                          self.dispatchcom1.year_create)

            context['a110_1'] = intern_utils.convert_to_vn_phone(self.dispatchcom1.phone_number)
            context['a112'] = intern_utils.convert_to_docx_string(self.dispatchcom1.address_en)

            mission = self.dispatchcom1.mission
            if '\n' in mission:
                pre = '<w:p><w:r><w:t>'
                post = '</w:t></w:r></w:p>'
                lineBreak = '<w:br/>'
                test = self.dispatchcom1.mission.replace('\n', lineBreak)
                mission = pre + test + post
            context['a128'] = intern_utils.convert_to_docx_string(mission)
            context['a129'] = self.dispatchcom1.number_of_employee
            context['a130'] = self.dispatchcom1.capital
            context['a131'] = self.dispatchcom1.revenue

            context['a40'] = intern_utils.date_time_in_jp(self.day_create_letter_promotion,
                                                          self.month_create_letter_promotion,
                                                          self.year_create_letter_promotion)

            tpl.render(context)
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    # @api.model
    # def search_read(self, domain=None, fields=None, offset=0, limit=None, order=None):
    #     if self.env.user.has_group('hh_intern.group_hs_user') and not self.env.user.has_group('hh_intern.group_hs_manager'):
    #         employee = self.env['hh.employee'].search([('user_id','=',self.env.user.id)])
    #         if employee:
    #             if domain is None:
    #                 domain = []
    #                 domain.append(('room_pttt','in',employee[0].department_hs.ids))
    #             else:
    #                 domain.append(('room_pttt', 'in', employee[0].department_hs.ids))
    #             return super(Invoice,self).search_read(domain,fields,offset,limit,order)
    #     else:
    #         return super(Invoice, self).search_read( domain, fields, offset, limit, order)


    status = fields.Selection([(4,'Khởi tạo'),(5,'Tiến cử'),(1,'Thi tuyển'),(2,'Chốt Trúng tuyển'),(3,'Hoàn thành'),
                               (6,'Tạm dừng'),(7,'Huỷ bỏ')],string='Trạng thái đơn hàng',default=4)

    @api.one
    def pass_exam(self):
        if self.status < 2:
            self.status = 2

    @api.one
    def finish_invoice(self):
        if self.status < 3:
            self.status = 3





    def create_certification_end_train(self,enterprise_id):
        doc_certification = self.env['intern.document'].search([('name', '=', "CCDT")], limit=1)
        if doc_certification:
            stream = BytesIO(doc_certification[0].attachment.decode("base64"))
            tpl = DocxTemplate(stream)
            tempFile = NamedTemporaryFile(delete=False)
            context = {}

            table_interns = []
            list_interns = self.interns_pass_doc
            if self.hoso_created:
                list_interns = self.interns_pass_doc_hs

            interns_pass = sorted(list_interns, key=lambda x: x.sequence_pass)
            for i, intern in enumerate(interns_pass):
                if intern.enterprise.id != int(enterprise_id):
                    continue
                info = {}
                info['name'] = intern.name_without_signal.upper()
                info['ns'] = intern_utils.date_time_in_jp(intern.day, intern.month, intern.year)
                table_interns.append(info)

            context['tbl_intern'] = table_interns
            context['a59']= intern_utils.convert_to_docx_string(self.guild.name_in_jp)
            context['a53'] = intern_utils.convert_to_docx_string(self.training_center.name_jp)
            context['a54'] = intern_utils.convert_to_docx_string(self.training_center.address_en)
            context['a56'] = intern_utils.convert_to_vn_phone(self.training_center.phone_number)
            context['a44'] = intern_utils.date_time_in_jp(self.day_start_training, self.month_start_training,
                                                          self.year_start_training)
            context['a45'] = intern_utils.date_time_in_jp(self.day_end_training, self.month_end_training,
                                                          self.year_end_training)

            context['a46'] = intern_utils.date_time_in_jp(self.day_create_plan_training_report_customer,
                                                          self.month_create_plan_training_report_customer,
                                                          self.year_create_plan_training_report_customer)
            context['a57'] = intern_utils.no_accent_vietnamese(self.training_center.responsive_person)

            tpl.render(context)
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile


    # order_tts = [6,0,()]
    #

    @api.model
    def create(self, vals):

        _logger.info("CREATE %s" % str(vals))
        result = super(Invoice, self).create(vals)


        # if 'dispatchcom2' in vals :
        #     self._cr.execute('DELETE FROM dispatch_order WHERE dispatch_order.invoice_id= %d'%result.id)
        #     if len(vals['dispatchcom2'])>0 and len(vals['dispatchcom2'][0][2]) > 0:
        #         query_insert = 'INSERT INTO dispatch_order (invoice_id, dispatch_id) VALUES (%d,ARRAY%s)'%(result.id,str(vals['dispatchcom2'][0][2]))
        #         self._cr.execute(query_insert)

        result['custom_id'] = 'ERP-%d' %result.id

        # user = self.env['res.users'].browse(self._uid)
        # if user.has_group('hh_intern.group_hs_user') and not user.has_group('hh_intern.group_tc_user'):
        #     result['hoso_created'] = True
        #     result['status'] = 2
        if 'hoso_created' in vals and vals['hoso_created']:
            result['status'] = 2
        return result



    @api.multi
    def write(self, vals):

        # _logger.info("WRITE1111 %s" % str(vals))

        # for intr in self.interns:
        #     _logger.info("Sequence %d"%intr.sequence)

        # if 'interns' in vals and 'intern_order' in vals:
        #
        #     def getKeySequence(item):
        #         return item[1]
        #     vals['intern_order'].sort(key=getKeySequence)
        #     arrayId = []
        #     for id in vals['intern_order']:
        #         arrayId.append(id[0])
        #     self._cr.execute('DELETE FROM intern_order WHERE intern_order.invoice_id = %d' % self.id)
        #     if len(arrayId) > 0:
        #         query_insert = 'INSERT INTO intern_order (invoice_id, intern_id) VALUES (%d,ARRAY%s)' % (
        #                     self.id, str(arrayId))
        #         self._cr.execute(query_insert)
        #     del vals['intern_order']
        #
        #
        # if 'interns_pass' in vals and 'intern_pass_order' in vals:
        #     def getKeySequence(item):
        #         return item[1]
        #
        #     arrayId = []
        #     vals['intern_pass_order'].sort(key=getKeySequence)
        #     for id in vals['intern_pass_order']:
        #         arrayId.append(id[0])
        #
        #     self._cr.execute('DELETE FROM internpass_order WHERE internpass_order.invoice_id = %d' % self.id)
        #     if len(arrayId) > 0:
        #         query_insert = 'INSERT INTO internpass_order (invoice_id, intern_id) VALUES (%d,ARRAY%s)' % (
        #             self.id, str(arrayId))
        #         self._cr.execute(query_insert)
        #     del vals['intern_pass_order']

        # if 'dispatchcom2' in vals:
        #     self._cr.execute('DELETE FROM dispatch_order WHERE dispatch_order.invoice_id= %d'%self.id)
        #
        #     if len(vals['dispatchcom2']) > 0 and len(vals['dispatchcom2'][0][2]) > 0:
        #         query_insert = 'INSERT INTO dispatch_order (invoice_id, dispatch_id) VALUES (%d,ARRAY%s)'%(
        #             self.id,str(vals['dispatchcom2'][0][2]))
        #         self._cr.execute(query_insert)
                # _logger.info("AAAA %s" % query_insert)
            # self._cr.execute('DELETE FROM dispatchcom2_intern_invoice_rel WHERE '
            #                  'dispatchcom2_intern_invoice_rel.intern_invoice_id = %d' % self.id)
            # if len(vals['dispatchcom2'][0][2]) > 0:
            #     query_insert = 'INSERT INTO ' \
            #                    'dispatchcom2_intern_invoice_rel (intern_invoice_id, dispatchcom2_id) VALUES '
            #     for i in vals['dispatchcom2'][0][2]:
            #         query_insert = query_insert + '(%d,%d),' % (self.id, i)
            #
            #     self._cr.execute(query_insert[:-1])
        pre_promoted = []
        if 'interns_clone' in vals:
            for intern in self.interns_clone:
                if intern.promoted:
                    pre_promoted.append(intern.intern_id.id)


        tmp = super(Invoice, self).write(vals)
        if 'interns_pass_doc' in vals or 'interns_pass_new' in vals:
            enterprises = []
            for intern in self.interns_pass_doc:
                if intern.enterprise and intern.enterprise.id not in enterprises:
                    enterprises.append(intern.enterprise.id)
            listtmp = sorted(self.interns_pass_doc, key=lambda x: x.sequence_pass)
            counter = 0
            for id in enterprises:
                for i, intern in enumerate(listtmp):
                    if not intern.enterprise:
                        intern.sequence_pass = len(self.interns_pass_doc)
                    elif intern.enterprise.id == id:
                        intern.sequence_pass = counter
                        counter += 1

        if 'interns_clone' in vals:
            last_promoted = []
            for intern in self.interns_clone:
                if self.enterprise_doc and not intern.enterprise:
                    intern.enterprise = self.enterprise_doc
                if intern.promoted:
                    last_promoted.append(intern.intern_id.id)

            sub_promoted = [item for item in pre_promoted if item not in last_promoted]
            current = []
            if self.promotion_removed:
                test = self.promotion_removed.split(',')
                for x in test:
                    current.append(int(x))

            current = current + sub_promoted
            lastest = [x for x in current if x not in last_promoted]
            self.promotion_removed = ','.join(map(str, lastest))

        return tmp

    promotion_removed = fields.Char('Promotion removed')

    # promotion_remove = fields.One2many('intern.removed','invoice_id','DS TTS bị loại tc')

    @api.multi
    def read(self, fields=None, load='_classic_read'):
        # _logger.info("CONTEXT %s"%self._context)
        result = super(Invoice,self).read(fields,load)
        # if len(result) == 1:
        #     for record in result:
                # if 'interns_pass' in record:
                #     try:
                #         self._cr.execute('SELECT intern_id FROM internpass_order WHERE internpass_order.invoice_id = %d'%record['id'])
                #         tmpresult = self._cr.dictfetchall()
                #         # _logger.info("RESULT %s"%tmpresult )
                #         if len(tmpresult)==1:
                #             ids = tmpresult[0]['intern_id']
                #             if len(ids) == len(record['interns_pass']):
                #                 record['interns_pass'] = ids
                #     except:
                #         _logger.info("Loi gi do")
                # if 'interns' in record:
                #     try:
                #         self._cr.execute('SELECT intern_id FROM intern_order WHERE intern_order.invoice_id = %d'%record['id'])
                #         tmpresult = self._cr.dictfetchall()
                #         # _logger.info("RESULT %s"%tmpresult )
                #         if len(tmpresult)==1:
                #             ids = tmpresult[0]['intern_id']
                #             if len(ids) == len(record['interns']):
                #                 record['interns'] = ids
                #     except:
                #         _logger.info("Loi gi do")

                # if 'dispatchcom2' in record:
                #     try:
                #         self._cr.execute(
                #             'SELECT dispatch_id FROM dispatch_order WHERE dispatch_order.invoice_id = %d' % record['id'])
                #         tmpresult2 = self._cr.dictfetchall()
                #         if len(tmpresult2) == 1:
                #             ids2 = tmpresult2[0]['dispatch_id']
                #             if len(ids2) == len(record['dispatchcom2']):
                #                 record['dispatchcom2'] = ids2
                #     except:
                #         _logger.info("LOI GI")
        return result

    @api.multi
    def unlink(self):
        # if not self:
        #     return True

        # if self.env.user.id == SUPERUSER_ID:
        #     for id in self.ids:
        #         # self._cr.execute('DELETE FROM internpass_order WHERE internpass_order.invoice_id = %s' % id)
        #         # self._cr.execute('DELETE FROM intern_order WHERE intern_order.invoice_id = %s' % id)
        #         self._cr.execute('DELETE FROM dispatch_order WHERE dispatch_order.invoice_id= %s' % id)
        #
        #     return super(Invoice, self).unlink()
        #
        # for id in self.ids:
        #     self._cr.execute('SELECT create_uid, name FROM intern_invoice WHERE intern_invoice.id = %s'%id)
        #     tmpresult = self._cr.fetchone()
        #     if tmpresult[0] != self.env.uid:
        #         raise ValidationError(u"Bạn không có quyền xoá đơn hàng %s"%tmpresult[1])

        # for id in self.ids:
            # self._cr.execute('DELETE FROM internpass_order WHERE internpass_order.invoice_id = %s'%id)
            # self._cr.execute('DELETE FROM intern_order WHERE intern_order.invoice_id = %s'%id)
            # self._cr.execute('DELETE FROM dispatch_order WHERE dispatch_order.invoice_id= %s'%id)

        return super(Invoice,self).unlink()


    #fields.Selection([(1,'Thi tuyển'),(2,'Trúng tuyển'),(3,'Hoàn thành'),(4,'Khởi tạo'),(5,'Tiến cử'),
    #                           (6,'Huỷ bỏ')],string='Trạng thái đơn hàng',default=4)
    @api.one
    def start_promotion(self):
        ensure_one = False
        for intern in self.interns_clone:
            if intern.promoted:
                ensure_one = True
                break

        if ensure_one:
            self.write({
                'status': 5,
            })
        else:
            raise ValidationError(u"Chưa có TTS nào trong danh sách tiến cử")

    @api.one
    def confirm_exam(self):
        ensure_one = False
        for intern in self.interns_clone:
            if intern.confirm_exam and not intern.issues_raise:
                ensure_one = True
                break
        if ensure_one:
            self.write({
                'status': 1,
            })
            for intern in self.interns_clone:
                intern.write({
                    'exam': True,
                })
        else:
            raise ValidationError(u"Chưa có TTS nào trong danh sách chốt thi tuyển")

    @api.one
    def confirm_pass(self):
        ensure_one = False
        if not self.date_join_school:
            raise ValidationError(u"Chưa có thông tin ngày nhập học trúng tuyển")
        if not self.date_pass:
            raise ValidationError(u"Chưa có thông tin ngày trúng tuyển")
        if not self.date_departure:
            raise ValidationError(u"Chưa có thông tin ngày dự kiến xuất cảnh")

        if not self.dispatchcom1:
            raise ValidationError(u"Chưa có thông tin pháp nhân")

        if not self.job_vi or not self.job_jp:
            raise ValidationError(u"Chưa có thông tin ngành nghề")

        for intern in self.interns_clone:
            if intern.pass_exam and not intern.issues_raise:
                ensure_one = True

                break
        if ensure_one:
            for intern in self.interns_clone:
                if intern.pass_exam and not intern.issues_raise:
                    if not intern.enterprise:
                        raise ValidationError(u"Chưa có thông tin xí nghiệp của TTS %s"%intern.name)
                    elif not intern.place_to_work:
                        raise ValidationError(u"Chưa có thông tin địa điểm làm việc của TTS %s" % intern.name)

            self.write({
                'status': 2,
            })
            for intern in self.interns_clone:
                intern.write({
                    'done_exam': True,
                    'sequence_pass': intern.sequence_exam
                })
        else:
            raise ValidationError(u"Chưa có TTS nào trong danh sách trúng tuyển")


    reason_pause_cancel = fields.Char('Lý do hoãn/huỷ đơn')
    date_pause_cancel_exam = fields.Datetime('Ngày hoãn/huỷ đơn')

    # @api.one
    def pause_invoice(self,reason):
        self.write({
            'previous_stt':self.status,
            'status': 6,
            'reason_pause_cancel':reason,
            'date_pause_cancel_exam':fields.datetime.today()
        })
        for intern in self.interns_clone:
            intern.write({
                'cancel_exam': True,
            })

    # @api.one
    def cancel_invoice(self,reason):
        self.write({
            'previous_stt': self.status,
            'status': 7,
            'reason_pause_cancel': reason,
            'date_pause_cancel_exam': fields.datetime.today()
        })
        for intern in self.interns_clone:
            intern.write({
                'cancel_exam': True,
            })

    previous_stt = fields.Integer('Trạng thái cũ')
    @api.one
    def revert_destroy(self):
        if self.status == 6 or self.status == 7:
            if self.previous_stt == 0:
                self.write({'status':4})
            else:
                self.write({'status': self.previous_stt})

    # @api.model
    # def get_issue_count(self,intern_id,invoice_id):
    #     self._cr.execute('SELECT * FROM intern_issue WHERE intern_id = %s AND invoice_id = %s' % (intern_id, invoice_id))
    #     dict = self._cr.dictfetchall()
    #     return dict
        # if tmpCount:
        #     return tmpCount
        # else:
        #     return 0
        # context = self._context.copy()
        # return {
        #     'name': 'form_name',
        #     'view_type': 'form',
        #     'view_mode': 'form',
        #     'res_model': 'intern.invoice',
        #     'view_id': False,
        #     'type': 'ir.actions.act_window',
        #     'target':'new',
        #     'context': context,
        # }


    date_confirm_form = fields.Date('Ngày chốt form')
    fee_policy = fields.Char('Phí+cơ chế đơn hàng')
    fee_departure = fields.Float('Phí xuất cảnh nam(USD)')
    fee_departure_women = fields.Float('Phí xuất cảnh nữ(USD)')
    fee_study = fields.Float('Tiền học(USD)')
    fee_eating = fields.Float('Tiền ăn(VND)')
    bonus_target = fields.Float('Thưởng nam')
    count_target = fields.Boolean('Tính chỉ tiêu mới nam')
    bonus_target_women = fields.Float('Thưởng nữ')
    count_target_women = fields.Boolean('Tính chỉ tiêu mới nữ')


    date_arrival_jp = fields.Date('Ngày tới Nhật')
    port_of_entry_jp =fields.Char('Cảng nhập cảnh')
    name_of_airline =fields.Char('Tên hãng hàng không')


    def create_visa_application_form(self,intern):
        doc_form = self.env['intern.document'].search([('name', '=', "VISA_FORM")], limit=1)
        if doc_form:
            stream = BytesIO(doc_form[0].attachment.decode("base64"))
            tpl = DocxTemplate(stream)
            tempFile = NamedTemporaryFile(delete=False)
            context = {}

            list_words = intern.name_without_signal.upper().split(' ')
            context['a3_1'] = list_words[list_words.length-2]
            context['a3_2'] = ' '.join(list_words[:-1])
            context['a7'] = intern_utils.date_time_in_en(intern.day,intern.month, intern.year)
            context['pro'] = intern_utils.no_accent_vietnamese(intern.province.name)

            if intern.gender == 'nam':
                context['a13_1'] = CheckedBox().init(24)
                context['a13_2'] = CheckBox().init(24)
            else:
                context['a13_1'] = CheckBox().init(24)
                context['a13_2'] = CheckedBox().init(24)

            if intern.marital_status.id is 0:
                context['a14_1'] = CheckedBox().init(24)
                context['a14_2'] = CheckBox().init(24)
                context['a14_3'] = CheckBox().init(24)
                context['a14_4'] = CheckBox().init(24)
            elif intern.marital_status.id is 1:
                context['a14_1'] = CheckBox().init(24)
                context['a14_2'] = CheckedBox().init(24)
                context['a14_3'] = CheckBox().init(24)
                context['a14_4'] = CheckBox().init(24)
            elif intern.marital_status.id is 2:
                context['a14_1'] = CheckBox().init(24)
                context['a14_2'] = CheckBox().init(24)
                context['a14_3'] = CheckBox().init(24)
                context['a14_4'] = CheckedBox().init(24)
            if intern.identity:
                context['a9'] = intern.identity
            else:
                context['a9'] = intern.identity_2

            if intern.passport_type == '0':
                context['a160_1'] = CheckedBox().init(24)
                context['a160_2'] = CheckBox().init(24)
                context['a160_3'] = CheckBox().init(24)
                context['a160_4'] = CheckBox().init(24)
            elif intern.passport_type == '1':
                context['a160_1'] = CheckBox().init(24)
                context['a160_2'] = CheckedBox().init(24)
                context['a160_3'] = CheckBox().init(24)
                context['a160_4'] = CheckBox().init(24)
            elif intern.passport_type == '2':
                context['a160_1'] = CheckBox().init(24)
                context['a160_2'] = CheckBox().init(24)
                context['a160_3'] = CheckedBox().init(24)
                context['a160_4'] = CheckBox().init(24)
            elif intern.passport_type == '3':
                context['a160_1'] = CheckBox().init(24)
                context['a160_2'] = CheckBox().init(24)
                context['a160_3'] = CheckBox().init(24)
                context['a160_4'] = CheckedBox().init(24)
            context['a161'] = intern.passport_no
            context['a162'] = intern.passport_place
            context['a163'] = intern.passport_date_issue
            context['a164'] = intern.passport_issuing_authority
            context['a165'] = intern.passport_date_expire

            context['a86'] = self.year_expire
            context['a166'] = self.date_arrival_jp.strftime('%d/%m/%Y')
            context['a167'] = self.port_of_entry_jp.upper()
            context['a168'] = self.name_of_airline.upper()

            context['a60'] = intern_utils.convert_to_docx_string(self.guild.name_in_en)
            context['a64'] = self.guild.phone_number
            context['a63'] = self.guild.post_code
            context['a62'] = intern_utils.convert_to_docx_string(self.guild.address_in_romaji)

            context['a169'] = intern.date_duration_previous_in_jp
            context['sdt'] = intern.phone_number

            tpl.render(context)
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_danh_sach_lao_dong(self,enterprise_id):
        doc_form = self.env['intern.document'].search([('name', '=', "DANH_SACH_LAO_DONG")], limit=1)
        if doc_form:
            stream = BytesIO(doc_form[0].attachment.decode("base64"))
            tpl = DocxTemplate(stream)
            # tempFile = NamedTemporaryFile(delete=False)
            tempFile = StringIO()
            context = {}

            table_interns = []
            list_interns = self.interns_pass_doc
            if self.hoso_created:
                list_interns = self.interns_pass_doc_hs

            interns_pass = sorted(list_interns, key=lambda x: x.sequence_pass)
            # interns_pass = sorted(self.interns_pass_doc, key=lambda x: x.sequence_pass)
            counter = 0
            for i, intern in enumerate(interns_pass):
                if intern.enterprise.id != int(enterprise_id):
                    continue
                counter+=1
                info = {}
                info['stt'] = counter
                info['name'] = intern.name.upper()
                info['ns'] = intern_utils.date_time_in_en(intern.day, intern.month, intern.year)
                if intern.gender == 'nu':
                    info['gender'] = u'Nữ'
                else:
                    info['gender'] = u'Nam'
                info['job'] = self.job_en
                info['departure'] = '%s/%s'%(self.month_departure_doc,self.year_departure_doc)
                tmps = intern.hktt.split(",")
                if len(tmps)>=3:
                    info['province'] = tmps[len(tmps)-1].strip()
                    info['district'] = tmps[len(tmps)-2].strip()
                    info['village'] = tmps[len(tmps)-3].strip()

                table_interns.append(info)

            context['tbl_intern'] = table_interns
            tpl.render(context)
            tpl.save(tempFile)
            tempFile.seek(0)
            # data = tempFile.read()
            # tempFile.close()
            return tempFile

    def create_check_list(self,enterprise_id):
        doc_form = self.env['intern.document'].search([('name', '=', "Checklist")], limit=1)
        if doc_form:
            stream = BytesIO(doc_form[0].attachment.decode("base64"))
            tpl = DocxTemplate(stream)
            tempFile = StringIO()
            context = {}
            enterprise = self.env['intern.enterprise'].browse(int(enterprise_id))
            context['xn'] = enterprise.name_jp
            context['nd'] = self.guild.name_in_jp
            counter = 0

            list_interns = self.interns_pass_doc
            if self.hoso_created:
                list_interns = self.interns_pass_doc_hs

            # interns_pass = sorted(list_interns, key=lambda x: x.sequence_pass)

            for intern in list_interns:
                if intern.enterprise and intern.enterprise.id == int(enterprise_id):
                    counter+=1
            context['sltts'] = counter

            tpl.render(context)
            tpl.save(tempFile)
            tempFile.seek(0)
            return tempFile


    def add_to_phieutraloi(self):
        phieutraloi = self.env['intern.phieutraloi'].search([('has_full','=',False)])
        count_man = 0
        count_women = 0
        for rec in phieutraloi:
            count_man+= rec.total_intern_men - rec.len_interns_man
            count_women+= rec.total_intern_women - rec.len_interns_women
        count_current_invoice_man = 0
        count_current_invoice_women = 0

        list_interns = self.interns_pass_doc
        if self.hoso_created:
            list_interns = self.interns_pass_doc_hs


        for intern in list_interns:
            if not intern.phieutraloi_id:
                if intern.gender and intern.gender=='nam':
                    count_current_invoice_man+=1
                else:
                    count_current_invoice_women+=1
        if count_current_invoice_man>count_man and count_current_invoice_women>count_women:
            raise ValidationError(u'Cần tạo phiếu trả lời mới, thiếu vị trí cho %d nam và %d nữ'%((count_current_invoice_man-count_man),(count_current_invoice_women-count_women)))
        elif count_current_invoice_man>count_man:
            raise ValidationError(u'Cần tạo phiếu trả lời mới, thiếu vị trí cho %d nam' % (count_current_invoice_man - count_man))
        elif count_current_invoice_women>count_women:
            raise ValidationError(u'Cần tạo phiếu trả lời mới, thiếu vị trí cho %d nữ' % (count_current_invoice_women - count_women))
        else:
            for intern in list_interns:
                for rec in phieutraloi:
                    if intern.gender and intern.gender == 'nam' and rec.total_intern_men>rec.len_interns_man:
                        rec.interns = [(4, intern.id)]
                        rec.len_interns_man +=1
                    elif intern.gender and intern.gender == 'nu' and rec.total_intern_women>rec.len_interns_women:
                        rec.interns = [(4, intern.id)]
                        rec.len_interns_women += 1

    custom_id = fields.Char('Mã tự động')

    custom_id_2 = fields.Char('Mã đơn hàng')

    # @api.multi
    # def _auto_invoice_id(self):
    #     for rec in self:
    #         rec.custom_id = 'ERP-%d'%(rec.id+1)

    targets = fields.One2many('intern.department','invoice_id',string=u'Khoán chỉ tiêu')

    type_recruitment = fields.Char('Hình thức tuyển dụng')

    room_td_care = fields.Many2one('department', string='Phòng TD')

    note_report = fields.Char('Ghi chú')

    date_join_school = fields.Date('Ngày nhập học trúng tuyển')

    date_pass = fields.Date('Ngày trúng tuyển')

    hoso_created = fields.Boolean('Đơn hàng của HS')
    # , default = lambda self: self._get_default_hoso_created()
    # @api.model
    # def _get_default_name(self):
    #     user = self.env['res.users'].browse(self._uid)
    #     if user.has_group('hh_intern.group_hs_user') and not user.has_group('hh_intern.group_tc_user'):
    #         return True
    #     return False

    def create_48(self,enterprise_id):
        doc_certification = self.env['intern.document'].search([('name', '=', "Doc4-8")], limit=1)
        if doc_certification:
            stream = BytesIO(doc_certification[0].attachment.decode("base64"))
            tpl = DocxTemplate(stream)
            tempFile = StringIO()
            context = {}
            table_dd = []
            date_start = datetime.strptime('%s/%s/%s'%(self.day_start_training,self.month_start_training,self.year_start_training),'%d/%m/%Y')
            for i in range(0,21):
                table_dd.append(date_start.strftime('%d/%m/%Y'))
                date_start = date_start+relativedelta(days=1)
            context['tbd'] = table_dd
            context['a44'] = intern_utils.date_time_in_jp(self.day_start_training, self.month_start_training,
                                                          self.year_start_training)
            context['a45'] = intern_utils.date_time_in_jp(self.day_end_training, self.month_end_training,
                                                          self.year_end_training)

            table_dd2 = []
            for i in range(0, 11):
                table_dd2.append(date_start.strftime('%d/%m/%Y'))
                date_start = date_start + relativedelta(days=1)
            context['tbd2'] = table_dd2
            context['a59'] = intern_utils.convert_to_docx_string(self.guild.name_in_jp)
            context['a68'] = self.guild.name_of_responsive_jp
            date_tmp = datetime.strptime('%s/%s/%s'%(self.day_end_training, self.month_end_training,
                                                          self.year_end_training),'%d/%m/%Y')

            date_tmp = date_tmp + relativedelta(days=1)
            day_tmp = '%02d'%date_tmp.day
            month_tmp = '%02d'%date_tmp.month
            context['a451'] = intern_utils.date_time_in_jp(day_tmp, month_tmp,'%d'%date_tmp.year)


            table_interns = []
            list_interns = None
            if self.interns_pass_doc and len(self.interns_pass_doc)>0:
                list_interns = sorted(self.interns_pass_doc, key=lambda x:x.sequence_pass)
            else:
                list_interns = sorted(self.interns_pass_doc_hs, key=lambda x: x.sequence_pass)

            date_departure_tmp = datetime.strptime('%s-%s-%s'%(self.year_departure_doc,self.month_departure_doc,self.day_departure_doc),'%Y-%m-%d')
            for i, intern in enumerate(list_interns):
                row = {}
                row['stt'] = '%d'%(i+1)
                row['name'] = u'%s'%intern_utils.no_accent_vietnamese(intern.name).upper()
                row['date_entry'] = u'%s'%intern_utils.date_time_in_jp('%02d'%date_departure_tmp.day,'%02d'%date_departure_tmp.month,'%d'%date_departure_tmp.year)
                table_interns.append(row)
            context['tb_interns'] = table_interns


            tpl.render(context)
            tpl.save(tempFile)
            tempFile.seek(0)
            return tempFile

    def create_1_27(self,enterprise_id):
        doc1_17_header = self.env['intern.document'].search([('name', '=', "Doc1-17-HEADER")], limit=1)
        doc1_17_footer = self.env['intern.document'].search([('name', '=', "Doc1-17-FOOTER")], limit=1)

        list_interns = None
        if not self.hoso_created:
            list_interns = sorted(self.interns_pass_doc, key=lambda x: x.sequence_pass)
        else:
            list_interns = sorted(self.interns_pass_doc_hs, key=lambda x: x.sequence_pass)

        if doc1_17_header and doc1_17_footer:
            stream = BytesIO(doc1_17_header[0].attachment.decode("base64"))
            target_document = Document(stream)

            stream2 = BytesIO(doc1_17_footer[0].attachment.decode("base64"))

            if len(list_interns)>1:
                tmp = len(list_interns) / 20
                for i in range(0, tmp ):
                    document = Document(stream2)
                    if i>0:
                        for table in document.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    if 'tbl_intern' in cell.text:
                                        cell.text = cell.text.replace('tbl_intern', 'tbl_intern%d' % i)
                                        break
                    # target_document.add_page_break()
                    for element in document.element.body:
                        target_document.element.body.append(element)
            tempFile1 = NamedTemporaryFile(delete=False)
            target_document.save(tempFile1.name)
            tempFile1.flush()
            tempFile1.close()
            tpl = DocxTemplate(tempFile1.name)

            context = {}
            context['date_promotion']= intern_utils.date_time_in_jp(self.day_create_letter_promotion,
                                                          self.month_create_letter_promotion,
                                                          self.year_create_letter_promotion)
            context['name'] = list_interns[0].name_without_signal.upper()
            context['name_jp'] = list_interns[0].name_in_japan.replace(u'・',' ')
            if len(list_interns)>1:
                context['total'] = u'（ほか　%d名）（別添名簿の通り）'%(len(list_interns)-1)
                for k in range(0, len(list_interns) / 20 + 1):
                    table_interns = []
                    for i in range(20 * k, 20 * k + 20):
                        if i >= len(list_interns):
                            break
                        intern = list_interns[i]
                        info = {}
                        info['stt'] = str(i + 1)
                        info['name'] = intern.name_without_signal.upper()
                        info['name_jp'] = intern.name_in_japan.replace(u'・',' ')
                        info['job_jp'] = self.job_jp
                        table_interns.append(info)

                    if k == 0:
                        context['tbl_intern'] = table_interns
                    else:
                        context['tbl_intern%d' % k] = table_interns
            tpl.render(context)

            tempFile = StringIO()
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.seek(0)
            if tempFile1 is not None:
                os.unlink(tempFile1.name)

            return tempFile

        return None


# class PromotionRemoved(models.Model):
#     _name = 'intern.removed'
#
#     invoice_id = fields.Many2one('intern.invoice', string='Đơn hàng')
#
#     intern_id = fields.Many2one('intern.intern', auto_join=True, string='TTS')
#     reason = fields.Char('Lý do')
