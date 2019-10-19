# -*- coding: utf-8 -*-
from odoo import models, fields, api

from tempfile import TemporaryFile, NamedTemporaryFile
from io import BytesIO
from datetime import datetime, date
from docxtpl import DocxTemplate, InlineImage, CheckedBox, CheckBox, RichText, Tick
from docx.shared import Mm, Inches
from docx import Document

# from openpyxl.styles.borders import Border, Side
# from openpyxl import load_workbook
from cStringIO import StringIO

import intern_utils
import os
from odoo.exceptions import UserError, ValidationError
import logging

_logger = logging.getLogger(__name__)
class InvoiceNew(models.Model):
    _inherit = 'intern.invoice'

    interns_promoted_doc = fields.One2many('intern.internclone', 'invoice_id',domain=[('promoted','=',True)])
    interns_exam_doc = fields.One2many('intern.internclone', 'invoice_id',domain=[('confirm_exam','=',True),('issues_raise','=',False)])
    interns_pass_doc = fields.One2many('intern.internclone', 'invoice_id',domain=[('pass_exam','=',True),('cancel_pass','=',False)])
    interns_pass_doc_hs = fields.One2many('intern.internclone', 'invoice_id_hs',domain=[('pass_exam','=',True),('cancel_pass','=',False)])


    @api.multi
    def start_translate_form(self):
        view_id = self.env.ref('hh_intern.view_doc_generate_clone').id
        context = self._context.copy()
        return {
            'name': 'form_name',
            'view_type': 'form',
            'view_mode': 'form',
            'res_model': 'intern.invoice',
            'view_id': view_id,
            'type': 'ir.actions.act_window',
            'res_id': self.id,
            'context': context,
        }

    @api.multi
    def create_doc_new(self):
        # _logger.info("CREATE DOC")
        # ensure_one = False
        # for intern in self.interns_exam_doc:
        #     if intern.confirm_exam and not intern.escape_exam:
        #         ensure_one = True
        #         break

        if self.interns_exam_doc is None or len(self.interns_exam_doc) is 0:
            raise ValidationError("Không có thực tập sinh nào trong danh sách thi tuyển")

        # for intern in self.interns_exam_doc:
            # if intern.date_of_birth is None:
            #     raise ValidationError(u"Thiếu thông tin %s của %s" % (u'ngày sinh', intern.name))
            #
            # if not intern.name_in_japan:
            #     raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Tên tiếng Nhật', intern.name))
            # if not intern.gender:
            #     raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Giới tính', intern.name))
            # if not intern.blood_group:
            #     raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Nhóm máu', intern.name))
            # if intern.vision_left is None or intern.vision_right is None:
            #     raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Thị lực', intern.name))
            # if not intern.iq_percentage or intern.iq_percentage is None or intern.iq_percentage is 0:
            #     raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Điểm IQ', intern.name))
            # if intern.check_kureperin is None:
            #     raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Điểm cộng dồn', intern.name))
            # if not intern.marital_status:
            #     raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Tình trạng hôn nhân', intern.name))
            # if not intern.province or not intern.address:
            #     raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Quê quán', intern.name))
            # if not intern.certification:
            #     raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Bằng cấp', intern.name))
            # if intern.certification.id > 2 and not intern.specialized:
            #     raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Chuyên ngành', intern.name))
            #
            # if intern.height is None:
            #     raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Chiều cao', intern.name))
            # if intern.weight is None:
            #     raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Cân nặng', intern.name))
            # if not intern.favourite or not intern.strong or \
            #         not intern.weak or not intern.family_income or not intern.motivation or \
            #         not intern.income_after_three_year or not intern.job_after_return or \
            #         not intern.prefer_object or not intern.memory or not intern.valuable:
            #     raise ValidationError(u"Thiếu thông tin %s của %s" % (u'bổ sung tiếng Nhật', intern.name))
            #
            # if intern.educations is None or len(intern.educations) is 0:
            #     raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Lý lịch Học tập', intern.name))
            # if intern.family_members is None or len(intern.family_members) is 0:
            #     raise ValidationError(u"Thiếu thông tin %s của %s" % (u'Thành viên gia đình', intern.name))

        _logger.info("CREATE DOC ENDDD")
        return {
            'type': 'ir.actions.act_url',
            'url': '/web/binary/download_document_new?model=intern.invoice&id=%s&filename=%s.zip' % (
            str(self.id), self.name),
            'target': 'self', }


    def createHeaderDocNew(self,gender=None):
        docs = self.env['intern.document'].search([('name', '=', "CV_HEAD")], limit=1)
        logo = self.env['intern.document'].search([('name', '=', "Logo")], limit=1)
        streamDoc = BytesIO(docs[0].attachment.decode("base64"))
        target_document = None
        tempFile1 = None
        lentmp = len(self.interns_exam_doc)
        if gender!=None:
            lentmp = 0
            for intern in self.interns_exam_doc:
                if intern.gender == gender:
                    lentmp+=1

        if lentmp>20:
            target_document = Document(streamDoc)
            tmp = lentmp/20
            for i in range(1,tmp+1):
                document = Document(streamDoc)
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if 'tbl_intern' in cell.text:
                                cell.text = cell.text.replace('tbl_intern','tbl_intern%d'%i)
                                break
                # target_document.add_page_break()
                for element in document.element.body:
                    target_document.element.body.append(element)

            tempFile1 = NamedTemporaryFile(delete=False)
            target_document.save(tempFile1.name)
            tempFile1.flush()
            tempFile1.close()

        if docs:
            tpl = None
            if target_document is None:
                stream = BytesIO(docs[0].attachment.decode("base64"))
                tpl = DocxTemplate(stream)
            else:
                tpl = DocxTemplate(tempFile1.name)

            context = {}
            interns = self.interns_exam_doc.sorted(key=lambda x: x.sequence_exam)

            counter_index = 0
            for k in range(0,lentmp/20+1):
                table_interns = []
                for i in range(20*k,20*k+20):
                    if i >= lentmp:
                        break
                    if gender == None:
                        intern = interns[i]
                    else:
                        while counter_index<len(interns):
                            if interns[counter_index].gender == gender:
                                intern = interns[counter_index]
                                counter_index += 1
                                break
                            counter_index+=1

                    info = {}
                    info['stt'] = str(i+1)
                    info['htk'] = intern.name_without_signal.upper()
                    info['htn'] = intern.name_in_japan
                    if intern.gender == 'nu':
                        info['gt'] = u'女'
                    else:
                        info['gt'] = u'男'
                    info['ns'] = intern_utils.date_time_in_jp(intern.day,intern.month,intern.year)
                    info['t'] = str(intern_utils.get_age_jp(datetime.now(),intern.day,intern.month,intern.year))
                    info['nm'] = intern.blood_group

                    left = 1.5 - (10.0 - intern.vision_left) / 10.0
                    right = 1.5 - (10.0 - intern.vision_right) / 10.0
                    info['tlt'] = "%.1f" % (left)
                    info['tlp'] = "%.1f" % (right)
                    if intern.preferred_hand == '0':
                        info['tt'] = u'右'
                    elif intern.preferred_hand == '1':
                        info['tt'] = u'左'
                    else:
                        info['tt'] = u'両手'
                    info['cc'] = str(intern.height)
                    info['cn'] = str(intern.weight)
                    info['iq'] = "%s%%"%intern.iq_percentage
                    info['ktk'] = intern.check_kureperin
                    info['hn'] = intern.marital_status.name_in_jp
                    if intern.province:
                        info['pro'] = intern_utils.no_accent_vietnamese(intern.province.name).upper()
                    info['bc'] = intern.certification.name_in_jp
                    table_interns.append(info)

                if k == 0:
                    context['tbl_intern'] = table_interns
                else:
                    context['tbl_intern%d'%k] = table_interns

            context['logo'] = InlineImage(tpl,BytesIO(logo[0].attachment.decode("base64")),width=Mm(35))
            if self.enterprise_doc:
                context['xn'] = self.enterprise_doc.name_jp
            context['nd'] = self.guild.name_in_jp
            context['now'] = intern_utils.date_time_in_jp(datetime.now().day,datetime.now().month,datetime.now().year)

            tpl.render(context)

            tempFile = NamedTemporaryFile(delete=False)
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()

            if tempFile1 is not None:
                os.unlink(tempFile1.name)

            return tempFile
        if tempFile1 is not None:
            os.unlink(tempFile1.name)
        return None


    @api.multi
    def create_list_exam_excel(self):
        return {
            'type': 'ir.actions.act_url',
            'url': '/web/binary/download_list_intern_exam_new?model=intern.invoice&id=%s&filename=%s.xlsx' % (
                str(self.id), self.name),
            'target': 'self', }


    # def create_list_exam_excel_data(self):
    #     docs = self.env['intern.document'].search([('name', '=', "CV_HEAD_LIST")], limit=1)
    #     streamDoc = BytesIO(docs[0].attachment.decode("base64"))
    #
    #     thin_border = Border(left=Side(style='thin'),
    #                          right=Side(style='thin'),
    #                          top=Side(style='thin'),
    #                          bottom=Side(style='thin'))
    #
    #     wb = load_workbook(streamDoc)
    #     sheet = wb.get_sheet_by_name(u'名簿リスト')
    #     sheet.cell(row=5, column=3).value = self.guild.name_in_jp
    #     sheet.cell(row=6, column=3).value = self.enterprise_doc.name_jp
    #     sheet.cell(row=6, column=17).value = intern_utils.date_time_in_jp(datetime.now().day,datetime.now().month,datetime.now().year)
    #     interns = self.interns_exam_doc.sorted(key=lambda x: x.sequence_exam)
    #     for i in range(len(self.interns_exam_doc)):
    #         intern = interns[i]
    #         sheet.cell(row=11+i, column=1).value = str(i + 1)
    #         sheet.cell(row=11+i, column=2).value = intern.name_without_signal.upper()
    #         sheet.cell(row=11+i, column=3).value = intern.name_in_japan
    #         if intern.gender == 'nu':
    #             sheet.cell(row=11+i, column=4).value = u'女'
    #         else:
    #             sheet.cell(row=11+i, column=4).value = u'男'
    #         sheet.cell(row=11+i, column=6).value = intern_utils.date_time_in_jp(intern.day, intern.month, intern.year)
    #         sheet.cell(row=11+i, column=7).value = str(intern_utils.get_age_jp(datetime.now(), intern.day, intern.month, intern.year))
    #         sheet.cell(row=11+i, column=8).value = intern.blood_group
    #
    #         left = 1.5 - (10.0 - intern.vision_left) / 10.0
    #         right = 1.5 - (10.0 - intern.vision_right) / 10.0
    #         sheet.cell(row=11+i, column=9).value = "%.1f" % (left)
    #         sheet.cell(row=11+i, column=10).value = "%.1f" % (right)
    #         if intern.preferred_hand == '0':
    #             sheet.cell(row=11+i, column=11).value = u'右'
    #         elif intern.preferred_hand == '1':
    #             sheet.cell(row=11+i, column=11).value = u'左'
    #         else:
    #             sheet.cell(row=11+i, column=11).value = u'両手'
    #         sheet.cell(row=11+i, column=12).value = str(intern.height)
    #         sheet.cell(row=11+i, column=13).value = str(intern.weight)
    #         sheet.cell(row=11+i, column=14).value = "%s%%" % intern.iq_percentage
    #         sheet.cell(row=11+i, column=15).value = intern.check_kureperin
    #         sheet.cell(row=11+i, column=16).value = intern.marital_status.name_in_jp
    #         sheet.cell(row=11+i, column=17).value = intern_utils.no_accent_vietnamese(intern.province.name).upper()
    #         sheet.cell(row=11+i, column=18).value = intern.certification.name_in_jp
    #
    #     for i in range(17):
    #         sheet.cell(row=8, column=i+1).border = thin_border
    #
    #     fp = StringIO()
    #     wb.save(fp)
    #     fp.seek(0)
    #     data = fp.read()
    #     fp.close()
    #     return data

        # workbook = xlsxwriter.Workbook(streamDoc)
        # worksheet = workbook.get_worksheet_by_name(u'名簿リスト')
        # worksheet.write('A2', 'Insert an image in a cell:')
        # workbook.close()


    date_receive_hard_profile = fields.Date('Ngày nhận hồ sơ cứng')
    date_receive_contract = fields.Date('Ngày nhận HĐL')
    date_send_letter_pro = fields.Date('Ngày trình thư')
    date_expected_send_to_customer = fields.Date('Ngày dự kiến gửi hồ sơ cho KH')
    date_real_send_to_customer = fields.Date('Ngày thực tế gửi hồ sơ cho KH')
    note_hs = fields.Char('Ghi chú')

    @api.multi
    def create_doc_new_2(self):
        _logger.info("create_doc_new_2")


    color_notice = fields.Char('Cảnh báo màu')

    @api.one
    def toggle_red(self):
        self.color_notice = '#e0596d'

    @api.one
    def toggle_yellow(self):
        self.color_notice = '#fff000'

    @api.one
    def toggle_green(self):
        self.color_notice = '#91c468'

    @api.one
    def toggle_remove_notice(self):
        self.color_notice = False

    @api.multi
    def create_doc_new_man(self):

        if self.interns_exam_doc is None or len(self.interns_exam_doc) is 0:
            raise ValidationError("Không có thực tập sinh nào trong danh sách thi tuyển")

        _logger.info("CREATE DOC ENDDD")
        return {
            'type': 'ir.actions.act_url',
            'url': '/web/binary/download_document_new?model=intern.invoice&id=%s&filename=%s.zip&gender=nam' % (
                str(self.id), self.name),
            'target': 'self', }

    @api.multi
    def create_doc_new_women(self):

        if self.interns_exam_doc is None or len(self.interns_exam_doc) is 0:
            raise ValidationError("Không có thực tập sinh nào trong danh sách thi tuyển")

        _logger.info("CREATE DOC ENDDD")
        return {
            'type': 'ir.actions.act_url',
            'url': '/web/binary/download_document_new?model=intern.invoice&id=%s&filename=%s.zip&gender=nu' % (
                str(self.id), self.name),
            'target': 'self', }

    @api.multi
    def read(self, fields=None, load='_classic_read'):
        if fields is not None:
            if 'interns_clone' in fields and 'interns_promoted' in fields and 'interns_confirm_exam' in fields:
                new_fields = list(fields)
                new_fields.remove('interns_promoted')
                new_fields.remove('interns_confirm_exam')
                new_fields.remove('interns_escape_exam')
                new_fields.remove('interns_pass_new')
                new_fields.remove('interns_preparatory')
                new_fields.remove('interns_cancel_pass')
                if 'interns_departure' in fields:
                    new_fields.remove('interns_departure')
                records = super(InvoiceNew, self).read(new_fields, load)
                for rec in records:
                    if 'interns_promoted' in fields:
                        rec['interns_promoted'] = rec['interns_clone']
                    if 'interns_confirm_exam' in fields:
                        rec['interns_confirm_exam'] = rec['interns_clone']
                    if 'interns_escape_exam' in fields:
                        rec['interns_escape_exam'] = rec['interns_clone']
                    if 'interns_pass_new' in fields:
                        rec['interns_pass_new'] = rec['interns_clone']
                    if 'interns_preparatory' in fields:
                        rec['interns_preparatory'] = rec['interns_clone']
                    if 'interns_cancel_pass' in fields:
                        rec['interns_cancel_pass'] = rec['interns_clone']
                    if 'interns_departure' in fields:
                        rec['interns_departure'] = rec['interns_clone']
                return records

        return super(InvoiceNew,self).read(fields,load)

