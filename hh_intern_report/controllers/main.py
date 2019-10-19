# -*- coding: utf-8 -*-
from odoo import http
from odoo.http import request
from odoo.addons.web.controllers.main import serialize_exception,content_disposition
from docx.enum.shape import WD_INLINE_SHAPE

from io import BytesIO, StringIO
from docxtpl import DocxTemplate, InlineImage, CheckedBox
from docx.shared import Mm, Inches, Pt
# from docx import Document
from tempfile import NamedTemporaryFile
import os

from odoo.addons.hh_intern.models import intern_utils
import zipfile

import logging
_logger = logging.getLogger(__name__)

class Document(http.Controller):

    @http.route('/document/preview/<model("intern.document"):document>', website=True)
    def preview(self,document, **kwargs):
        """
        Basic Hello World example
        """
        attachment_ids = request.env['ir.attachment'].search(
            [('res_model', '=', 'intern.document'), ('res_id', '=', document.id)])

        for attach in attachment_ids:
            url = 'http://localhost:8069' + '/web/content/' + str(attach.id)+"/temp.pdf&embedded=true"
            return request.render('hh_intern_report.preview', {'data': url})

    @http.route('/document/createprint/<model("intern.document"):document>', website=True)
    def createprint(self,document, **kwargs):
        return request.render(
            'hh_intern_report.add')

    @http.route('/document/add', website=True)
    def add(self, **kwargs):
        # interns = request.env['intern.intern'].search()
        return request.render(
            'hh_intern_report.add')

    @http.route('/document/edit/<model("intern.document"):document>', website=True)
    def edit(self, document, **kwargs):
        attachment_ids = request.env['ir.attachment'].search(
            [('res_model', '=', 'intern.document'), ('res_id', '=', document.id)])
        attachment_name = "";
        if attachment_ids is not None:
            attachment_name = attachment_ids[0].datas_fname
        return request.render(
            'hh_intern_report.edit', {'attachment_name': attachment_name, 'document': document})

    @http.route('/document', website=True)
    def index(self, **kwargs):
        """
        Todo list page
        """
        DocTask = request.env['intern.document']
        tasks = DocTask.search([])
        return request.render(
            'hh_intern_report.index', {'documents': tasks})

    # @http.route('/web/binary/download_document', type='http', auth="public")
    # def download_cv(self, model,id, filename = None, **kwargs):
    #     invoice = request.env[model].search([('id', '=',id)])
    #     document = request.env['intern.document'].browse(invoice.document.id)
    #     # _logger.info(vals['interns'])
    #
    #     finalDoc = invoice.createHeaderDoc(invoice.interns)
    #
    #     for i,intern in enumerate(invoice.interns):
    #         _logger.info("Intern ")
    #         childDoc = invoice.createCVDoc(document, intern,i)
    #
    #         if childDoc is not None:
    #             # _logger.info("xxx %d" % len(childDoc.tables))
    #
    #             for table in childDoc.tables:
    #                 _logger.info("XAAAAAAAA %d" % len(table.rows))
    #                 for row in table.rows:
    #                     _logger.info("xxxxxx %d" % len(row.cells))
    #                     for cell in row.cells:
    #                         _logger.info("DDDDD %d" % len(cell.part._package.image_parts))
    #
    #
    #             finalDoc.add_page_break()
    #             for element in childDoc.element.body:
    #                 finalDoc.element.body.append(element)
    #
    #             _logger.info("AAAA %d" % len(childDoc.part._package.image_parts))
    #             _logger.info("BBB %d" % len(finalDoc.part._package.image_parts))
    #             _logger.info("EEEE %d" % len(finalDoc.inline_shapes))
    #
    #             # for image in childDoc.part._package.image_parts:
    #             #     image_bytes = image.blob
    #             #     image_stream = BytesIO(image_bytes)
    #             #     rId, imageNew = finalDoc.part.get_or_add_image(image_stream)
    #
    #             iterate = iter(childDoc.part._package.image_parts)
    #             for i, shape in enumerate(finalDoc.inline_shapes):
    #                 if shape.type == WD_INLINE_SHAPE.PICTURE:
    #                     _logger.info('i= %d, Shape is an embedded picture' % (i))
    #
    #                     assert shape.type == WD_INLINE_SHAPE.PICTURE
    #                     inline = shape._inline
    #                     rId = inline.xpath('./a:graphic/a:graphicData/pic:pic/pic:blipFill/a:blip/@r:embed')[0]
    #                     image_part = childDoc.part.related_parts[rId]
    #                     # image_bytes = image_part.blob
    #
    #                     imageSource = iterate.next()
    #                     image_stream = BytesIO(imageSource.blob)
    #                     inlineNew = finalDoc.part.new_pic_inline(image_stream,image_part.default_cx,image_part.default_cy)
    #                     shape._inline = inlineNew
    #
    #                     # # write the image bytes to a file (or BytesIO stream) and feed it to document.add_picture(), maybe:
    #                     # image_stream = BytesIO(image_bytes)
    #                     # finalDoc.add_picture(image_stream)
    #                 elif shape.type == WD_INLINE_SHAPE.LINKED_PICTURE:
    #                     _logger.info( 'i= %d, Shape is a picture, but actual image file is not in this package' % (i))
    #                 else:
    #                     _logger.info( 'i= %d, Shape is not a picture, got %s' % (i, shape.type))
    #
    #     # finalDoc.save("ttt.docx")
    #     byteIo = BytesIO()
    #     finalDoc.save(byteIo)
    #     byteIo.seek(0)
    #     return request.make_response(byteIo,
    #                                  [('Content-Type', 'application/octet-stream'),
    #                                   ('Content-Disposition', content_disposition(filename))])

    # @http.route('/web/binary/download_document', type='http', auth="public")
    # def download_cv(self, model, id, filename=None, **kwargs):
    #     invoice = request.env[model].search([('id', '=', id)])
    #     document = request.env['intern.document'].browse(invoice.document.id)
    #     # _logger.info(vals['interns'])
    #
    #     finalDoc = invoice.createHeaderDoc(invoice.interns)
    #
    #     # for i, intern in enumerate(invoice.interns):
    #     #     childDoc = invoice.createCVDoc(document, intern, i)
    #     #     if childDoc is not None:
    #     #         finalDoc.add_page_break()
    #     #         for element in childDoc.element.body:
    #     #             finalDoc.element.body.append(element)
    #
    #     byteIo = BytesIO()
    #     finalDoc.save(byteIo)
    #     byteIo.seek(0)
    #
    #     tpl = DocxTemplate(byteIo)
    #
    #     context = {}
    #     for i,intern in enumerate(invoice.interns):
    #         streamAvatar = None
    #         if intern.avatar is not None:
    #             streamAvatar = BytesIO(intern.avatar.decode("base64"))
    #             if streamAvatar is not None:
    #                 context['avatar%d'%i]= InlineImage(tpl, streamAvatar, width=Mm(20))
    #
    #     context['co_mu_mau'] = CheckedBox()
    #
    #     tpl.render(context)
    #
    #     byteIoRespond = BytesIO()
    #     tpl.save(byteIoRespond)
    #     byteIoRespond.seek(0)
    #     # f = NamedTemporaryFile()
    #     # tpl.save(f.name)
    #     # f.seek(0)
    #     return request.make_response(byteIoRespond,
    #                                  [('Content-Type', 'application/octet-stream'),
    #                                   ('Content-Disposition', content_disposition(filename))])

        # tpl.save("test111.docx")

    # @http.route('/web/binary/test_doc', type='http', auth="public")
    # def test_doc(self,**kwargs):
    #     document = request.env['intern.document'].search([('name', '=', 'CV')], limit=1)
    #     streamDoc = BytesIO(document[0].attachment.decode("base64"))
    #     target_document = Document(streamDoc)
    #     for i in range(0,3):
    #         tmp = Document(streamDoc)
    #         for element in tmp.element.body:
    #             target_document.element.body.append(element)
    #
    #     byteIoRespond = BytesIO()
    #     target_document.save(byteIoRespond)
    #     byteIoRespond.seek(0)
    #     return request.make_response(byteIoRespond,
    #                                             [('Content-Type', 'application/octet-stream'),
    #                                              ('Content-Disposition', "test.docx")])


    @http.route('/web/binary/download_document', type='http', auth="public")
    def download_cv(self, model, id, filename=None, **kwargs):
        _logger.info("DOWNLOAD CV")
        invoice = request.env[model].search([('id', '=', id)])
        document = request.env['intern.document'].search([('name', '=', 'CV')], limit=1)
        finalDoc = invoice.createHeaderDoc()
        # byteIo = BytesIO()
        # finalDoc.save(byteIo)
        # byteIo.seek(0)
        reponds = BytesIO()
        archive = zipfile.ZipFile(reponds, 'w', zipfile.ZIP_DEFLATED)
        # merge_doc = None
        if finalDoc is not None:
            archive.write(finalDoc.name,u"名簿リスト.docx")

            # merge_doc = Document(finalDoc.name)
            os.unlink(finalDoc.name)
        else:
            return

        if invoice.order:
            ids = []
            for intern in invoice.interns:
                ids.append(intern.id)
            list = None
            try:
                list = request.env['intern.intern'].search([('id', 'in', ids)], order="%s" % (invoice.order))
            except:
                list = invoice.interns

            for i, intern in enumerate(list):
                childDoc = invoice.createCVDoc(document[0], intern, i)
                archive.write(childDoc.name,'cv_%d_%s.docx'%((i+1),intern_utils.name_with_underscore(intern.name)))

                # tmpDoc = Document(childDoc.name)
                # for element in tmpDoc.element.body:
                #     merge_doc.element.body.append(element)

                os.unlink(childDoc.name)

        else:
            for i, intern in enumerate(invoice.interns):
                childDoc = invoice.createCVDoc(document[0], intern, i)
                archive.write(childDoc.name,'cv_%d_%s.docx'%((i+1),intern_utils.name_with_underscore(intern.name)))

                # tmpDoc = Document(childDoc.name)
                # for element in tmpDoc.element.body:
                #     merge_doc.element.body.append(element)

                os.unlink(childDoc.name)

        # tempFile = NamedTemporaryFile(delete=False)
        # merge_doc.save(tempFile.name)
        # archive.write(tempFile.name,"full.docx")
        # os.unlink(tempFile.name)

        archive.close()
        reponds.flush()
        ret_zip = reponds.getvalue()
        reponds.close()

        return request.make_response(ret_zip,
                                         [('Content-Type', 'application/zip'),
                                          ('Content-Disposition', content_disposition(filename))])

    @http.route('/web/binary/download_proletter_document', type='http', auth="public")
    def download_proletter_document(self, model, id, filename, **kwargs):
        invoice = request.env[model].search([('id', '=', id)])

        reponds = BytesIO()
        archive = zipfile.ZipFile(reponds, 'w', zipfile.ZIP_DEFLATED)

        doc_list_send = invoice.create_list_of_sent_en()
        archive.write(doc_list_send.name, u'推薦書 - ENG.docx')
        os.unlink(doc_list_send.name)

        doc_list_send_jp = invoice.create_list_of_sent_jp()
        archive.write(doc_list_send_jp.name, u'推薦書.docx')
        os.unlink(doc_list_send_jp.name)
        archive.close()
        reponds.flush()
        ret_zip = reponds.getvalue()
        reponds.close()

        return request.make_response(ret_zip,
                                     [('Content-Type', 'application/zip'),
                                      ('Content-Disposition', content_disposition(filename))])

    @http.route('/web/binary/download_extern_document_specific', type='http', auth="public")
    def download_extern_document_specific(self, model, id, document,filename=None, **kwargs):
        invoice = request.env[model].search([('id', '=', id)])

        reponds = BytesIO()
        archive = zipfile.ZipFile(reponds, 'w', zipfile.ZIP_DEFLATED)

        if document == 'Doc1-3':
            for i, intern in enumerate(invoice.interns_pass):
                doc1_3 = invoice.create_doc_1_3(intern, i)
                archive.write(doc1_3.name, '1_3_%d_%s.docx' % ((i + 1), intern_utils.name_with_underscore(intern.name)))
                os.unlink(doc1_3.name)
        elif document == 'Doc1-10':
            for i, intern in enumerate(invoice.interns_pass):
                doc1_10 = invoice.create_doc_1_10(intern)
                archive.write(doc1_10.name,
                              '1_10_%d_%s.docx' % ((i + 1), intern_utils.name_with_underscore(intern.name)))
                os.unlink(doc1_10.name)
        elif document == 'Doc1-13':
            doc1_13_1 = invoice.create_1_13_1()

            archive.write(doc1_13_1.name, u'1-13号 HOANG HUNG JAPAN 訓連センター.docx')
            os.unlink(doc1_13_1.name)

            doc1_13_2 = invoice.create_1_13_2()
            archive.write(doc1_13_2.name, u'1-13号HOANG HUNG 会社.docx')
            os.unlink(doc1_13_2.name)
        elif document == 'Doc1-20':
            doc1_20 = invoice.create_doc_1_20()
            archive.write(doc1_20.name, '1_20.docx')
            os.unlink(doc1_20.name)
        elif document == 'Doc1-21':
            for i, intern in enumerate(invoice.interns_pass):
                doc1_21 = invoice.create_doc_1_21(intern)
                archive.write(doc1_21.name, '1_21_%d_%s.docx' % ((i + 1), intern_utils.name_with_underscore(intern.name)))
                os.unlink(doc1_21.name)
        elif document == 'Doc1-28':
            for i, intern in enumerate(invoice.interns_pass):
                doc1_28 = invoice.create_doc_1_28(intern,i)
                archive.write(doc1_28.name, '1_28_%d_%s.docx' % ((i + 1), intern_utils.name_with_underscore(intern.name)))
                os.unlink(doc1_28.name)
        elif document == 'Doc1-29':
            doc1_29 = invoice.create_doc_1_29()
            archive.write(doc1_29.name, '1_29.docx')
            os.unlink(doc1_29.name)
        elif document == 'HDPC':
            for i, intern in enumerate(invoice.interns_pass):
                hdtn = invoice.create_hdtn(intern)
                archive.write(hdtn.name, 'hdtn_%d_%s.docx' % ((i + 1), intern_utils.name_with_underscore(intern.name)))
                os.unlink(hdtn.name)

                hdtv = invoice.create_hdtv(intern)
                archive.write(hdtv.name, 'hdtv_%d_%s.docx' % ((i + 1), intern_utils.name_with_underscore(intern.name)))
                os.unlink(hdtv.name)
        elif document == 'PROLETTER':
            doc_list_send = invoice.create_list_of_sent_en()
            archive.write(doc_list_send.name, u'推薦書 - ENG.docx')
            os.unlink(doc_list_send.name)

            doc_list_send_jp = invoice.create_list_of_sent_jp()
            archive.write(doc_list_send_jp.name, u'推薦書.docx')
            os.unlink(doc_list_send_jp.name)


        archive.close()
        reponds.flush()
        ret_zip = reponds.getvalue()
        reponds.close()

        return request.make_response(ret_zip,
                                     [('Content-Type', 'application/zip'),
                                      ('Content-Disposition', content_disposition(filename))])

    @http.route('/web/binary/download_extern_document', type='http', auth="public")
    def download_extern_document(self, model, id, filename=None, **kwargs):
        invoice = request.env[model].search([('id', '=', id)])

        reponds = BytesIO()
        archive = zipfile.ZipFile(reponds, 'w', zipfile.ZIP_DEFLATED)

        checklist = request.env['intern.document'].search([('name', '=', "Checklist")], limit=1)
        if checklist:
            stream = BytesIO(checklist[0].attachment.decode("base64"))
            tpl = DocxTemplate(stream)
            tempFile = NamedTemporaryFile(delete=False)
            tpl.render({})
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()
            archive.write(tempFile.name, 'Checklist.docx')
            os.unlink(tempFile.name)

        doc1_13_1 = invoice.create_1_13_1()

        archive.write(doc1_13_1.name, u'1-13号 HOANG HUNG JAPAN 訓連センター.docx')
        os.unlink(doc1_13_1.name)

        doc1_13_2 = invoice.create_1_13_2()
        archive.write(doc1_13_2.name, u'1-13号HOANG HUNG 会社.docx')
        os.unlink(doc1_13_2.name)

        master = invoice.create_master()
        archive.write(master.name, 'Master.docx')
        os.unlink(master.name)

        doc1_29 = invoice.create_doc_1_29()
        archive.write(doc1_29.name, '1_29.docx')
        os.unlink(doc1_29.name)

        doc_list_send = invoice.create_list_of_sent_en()
        archive.write(doc_list_send.name, u'推薦書 - ENG.docx')
        os.unlink(doc_list_send.name)

        doc_list_send_jp = invoice.create_list_of_sent_jp()
        archive.write(doc_list_send_jp.name, u'推薦書.docx')
        os.unlink(doc_list_send_jp.name)

        doc1_20 = invoice.create_doc_1_20()
        archive.write(doc1_20.name, '1_20.docx')
        os.unlink(doc1_20.name)

        for i, intern in enumerate(invoice.interns_pass):
            doc1_3 = invoice.create_doc_1_3(intern, i)
            archive.write(doc1_3.name, '1_3_%d_%s.docx' % ((i+1),intern_utils.name_with_underscore(intern.name)))
            os.unlink(doc1_3.name)

            doc1_10 = invoice.create_doc_1_10(intern)
            archive.write(doc1_10.name, '1_10_%d_%s.docx' % ((i+1),intern_utils.name_with_underscore(intern.name)))
            os.unlink(doc1_10.name)



            doc1_21 = invoice.create_doc_1_21(intern)
            archive.write(doc1_21.name, '1_21_%d_%s.docx' % ((i+1),intern_utils.name_with_underscore(intern.name)))
            os.unlink(doc1_21.name)

            doc1_28 = invoice.create_doc_1_28(intern,i)
            archive.write(doc1_28.name, '1_28_%d_%s.docx' % ((i+1),intern_utils.name_with_underscore(intern.name)))
            os.unlink(doc1_28.name)

            hdtn = invoice.create_hdtn(intern)
            archive.write(hdtn.name, 'hdtn_%d_%s.docx' % ((i+1),intern_utils.name_with_underscore(intern.name)))
            os.unlink(hdtn.name)

            hdtv = invoice.create_hdtv(intern)
            archive.write(hdtv.name, 'hdtv_%d_%s.docx' % ((i+1),intern_utils.name_with_underscore(intern.name)))
            os.unlink(hdtv.name)

        archive.close()
        reponds.flush()
        ret_zip = reponds.getvalue()
        reponds.close()

        return request.make_response(ret_zip,
                                     [('Content-Type', 'application/zip'),
                                      ('Content-Disposition', content_disposition(filename))])
